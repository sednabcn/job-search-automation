#!/usr/bin/env python3
"""
Job Scoring System
Evaluates job postings against CV and preferences
Supports PDF, DOCX, and TXT file formats
"""

import json
import re
from pathlib import Path
from datetime import datetime, timedelta
from typing import Dict, List, Set
import sys

class JobScorer:
    def __init__(self, cv_path: str, preferences_path: str = None):
        """
        Initialize job scorer with CV and preferences
        
        Args:
            cv_path: Path to CV file (pdf, docx, or txt)
            preferences_path: Optional path to scoring preferences JSON
        """
        self.cv_keywords = self.extract_cv_keywords(cv_path)
        self.preferences = self.load_preferences(preferences_path)
    
    def extract_cv_keywords(self, cv_path: str) -> Set[str]:
        """Extract technical skills and keywords from CV (supports PDF, DOCX, TXT)"""
        cv_file = Path(cv_path)
        
        if not cv_file.exists():
            print(f"⚠️  CV file not found: {cv_path}")
            return set()
        
        # Get file extension
        ext = cv_file.suffix.lower()
        
        try:
            # Extract text based on file type
            if ext == '.pdf':
                content = self._extract_pdf_text(cv_file)
            elif ext in ['.docx', '.doc']:
                content = self._extract_docx_text(cv_file)
            elif ext == '.txt':
                content = cv_file.read_text(encoding='utf-8')
            else:
                print(f"⚠️  Unsupported CV format: {ext}")
                print(f"   Supported formats: .pdf, .docx, .txt")
                return set()
            
            content = content.lower()
            
            # Common technical keywords to extract
            tech_patterns = [
                r'\b(python|java|javascript|typescript|go|rust|c\+\+|c#|php|ruby|swift|kotlin)\b',
                r'\b(react|angular|vue|django|flask|spring|node\.?js|express|fastapi)\b',
                r'\b(aws|azure|gcp|google cloud|docker|kubernetes|terraform|ansible)\b',
                r'\b(sql|nosql|mongodb|postgresql|mysql|redis|cassandra|dynamodb)\b',
                r'\b(machine learning|ml|ai|data science|nlp|deep learning|tensorflow|pytorch)\b',
                r'\b(git|ci/cd|jenkins|github actions|gitlab|circleci|travis)\b',
                r'\b(agile|scrum|kanban|jira|confluence)\b',
                r'\b(rest|api|graphql|microservices|serverless)\b',
                r'\b(linux|unix|bash|shell|windows|macos)\b',
            ]
            
            keywords = set()
            for pattern in tech_patterns:
                matches = re.findall(pattern, content, re.IGNORECASE)
                keywords.update([m.lower() for m in matches])
            
            # Also extract common skills (simple word extraction)
            words = re.findall(r'\b[a-z]{3,}\b', content)
            
            # Filter for technical terms
            technical_terms = {
                'frontend', 'backend', 'fullstack', 'devops', 'cloud', 'database',
                'security', 'testing', 'deployment', 'automation', 'optimization',
                'scalability', 'performance', 'architecture', 'design', 'development'
            }
            
            for word in words:
                if word in technical_terms:
                    keywords.add(word)
            
            print(f"✅ Extracted {len(keywords)} keywords from CV ({ext} format)")
            if keywords:
                print(f"   Sample keywords: {', '.join(list(keywords)[:10])}")
            
            return keywords
            
        except Exception as e:
            print(f"⚠️  Error extracting CV keywords: {e}")
            print(f"   Falling back to empty keyword set")
            return set()
    
    def _extract_pdf_text(self, pdf_path: Path) -> str:
        """Extract text from PDF file"""
        try:
            import PyPDF2
            
            text = ""
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                print(f"   📄 Reading PDF: {len(pdf_reader.pages)} pages")
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    text += page_text + "\n"
                    
                    if page_num == 0 and not page_text.strip():
                        print(f"   ⚠️  Warning: Page {page_num + 1} appears empty")
            
            if not text.strip():
                raise ValueError("PDF appears to be empty or image-based (scanned PDF)")
            
            print(f"   ✅ Extracted {len(text)} characters from PDF")
            return text
            
        except ImportError:
            print("   ❌ PyPDF2 not installed")
            print("   Install with: pip install PyPDF2")
            raise ImportError(
                "PyPDF2 is required to read PDF files. "
                "Install it with: pip install PyPDF2"
            )
        except Exception as e:
            print(f"   ❌ Error reading PDF: {e}")
            raise
    
    def _extract_docx_text(self, docx_path: Path) -> str:
        """Extract text from Word document"""
        try:
            import docx
            
            doc = docx.Document(docx_path)
            
            # Extract paragraphs
            paragraphs = [paragraph.text for paragraph in doc.paragraphs]
            text = "\n".join(paragraphs)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += "\n" + cell.text
            
            if not text.strip():
                raise ValueError("Word document appears to be empty")
            
            print(f"   ✅ Extracted {len(text)} characters from DOCX")
            return text
            
        except ImportError:
            print("   ❌ python-docx not installed")
            print("   Install with: pip install python-docx")
            raise ImportError(
                "python-docx is required to read Word files. "
                "Install it with: pip install python-docx"
            )
        except Exception as e:
            print(f"   ❌ Error reading DOCX: {e}")
            raise
    
    def load_preferences(self, preferences_path: str = None) -> Dict:
        """Load scoring preferences or use defaults"""
        default_prefs = {
            'required_skills': [],
            'preferred_skills': [],
            'salary_min': 0,
            'salary_max': 999999,
            'locations': ['remote', 'london', 'uk'],
            'experience_level': 'mid',  # junior, mid, senior
            'company_preferences': [],
            'recency_days': 30,
            'weights': {
                'keywords': 40,
                'salary': 20,
                'location': 15,
                'company': 10,
                'experience': 10,
                'recency': 5
            }
        }
        
        if preferences_path and Path(preferences_path).exists():
            try:
                with open(preferences_path, 'r') as f:
                    custom_prefs = json.load(f)
                    default_prefs.update(custom_prefs)
                print(f"✅ Loaded custom preferences from: {preferences_path}")
            except Exception as e:
                print(f"⚠️  Error loading preferences: {e}")
                print(f"   Using default preferences")
        else:
            print(f"ℹ️  Using default preferences")
        
        return default_prefs
    
    def score_job(self, job: Dict) -> Dict:
        """
        Score a job posting (0-100)
        
        Args:
            job: Job dictionary with fields like title, description, company, etc.
        
        Returns:
            Dictionary with score and breakdown
        """
        scores = {
            'keywords': self._score_keywords(job),
            'salary': self._score_salary(job),
            'location': self._score_location(job),
            'company': self._score_company(job),
            'experience': self._score_experience(job),
            'recency': self._score_recency(job)
        }
        
        # Calculate weighted total
        weights = self.preferences['weights']
        total_score = sum(scores[key] * weights[key] / 100 for key in scores)
        
        return {
            'job_id': job.get('id', job.get('job_id', 'unknown')),
            'job_title': job.get('title', 'Unknown'),
            'company': job.get('company', 'Unknown'),
            'total_score': round(total_score, 1),
            'match_percentage': round(total_score, 1),
            'breakdown': scores,
            'recommendation': self._get_recommendation(total_score),
            'scored_at': datetime.now().isoformat()
        }
    
    def _score_keywords(self, job: Dict) -> float:
        """Score based on keyword match (0-100)"""
        description = job.get('description', '').lower()
        title = job.get('title', '').lower()
        requirements = job.get('requirements', '').lower()
        
        full_text = f"{title} {description} {requirements}"
        
        # Count keyword matches
        matches = sum(1 for kw in self.cv_keywords if kw in full_text)
        
        if not self.cv_keywords:
            return 50.0  # Neutral if no CV keywords
        
        match_ratio = matches / len(self.cv_keywords)
        return min(match_ratio * 100, 100)
    
    def _score_salary(self, job: Dict) -> float:
        """Score based on salary alignment (0-100)"""
        salary = job.get('salary', '')
        salary_min = job.get('salary_min', 0)
        salary_max = job.get('salary_max', 0)
        
        pref_min = self.preferences['salary_min']
        pref_max = self.preferences['salary_max']
        
        # If no salary info, neutral score
        if not salary and not salary_min and not salary_max:
            return 60.0
        
        # Parse salary if string
        if isinstance(salary, str):
            numbers = re.findall(r'\d+(?:,\d{3})*(?:\.\d+)?', salary.replace(',', ''))
            if numbers:
                salary_min = float(numbers[0])
                salary_max = float(numbers[-1]) if len(numbers) > 1 else salary_min
        
        # Check if salary meets minimum
        if salary_max and salary_max < pref_min:
            return 0.0
        
        # Full score if within range
        if salary_min >= pref_min and salary_max <= pref_max:
            return 100.0
        
        # Partial score if overlaps
        if salary_min <= pref_max and salary_max >= pref_min:
            return 75.0
        
        return 50.0
    
    def _score_location(self, job: Dict) -> float:
        """Score based on location preference (0-100)"""
        location = job.get('location', '').lower()
        
        if not location:
            return 50.0
        
        pref_locations = [loc.lower() for loc in self.preferences['locations']]
        
        # Check for matches
        for pref_loc in pref_locations:
            if pref_loc in location:
                return 100.0
        
        # Partial match for general terms
        if any(term in location for term in ['remote', 'hybrid', 'work from home']):
            return 80.0
        
        return 30.0
    
    def _score_company(self, job: Dict) -> float:
        """Score based on company preference (0-100)"""
        company = job.get('company', '').lower()
        
        if not company:
            return 50.0
        
        company_prefs = [c.lower() for c in self.preferences.get('company_preferences', [])]
        
        if not company_prefs:
            return 50.0  # Neutral if no preferences
        
        # Check if preferred company
        if any(pref in company for pref in company_prefs):
            return 100.0
        
        return 40.0
    
    def _score_experience(self, job: Dict) -> float:
        """Score based on experience level match (0-100)"""
        title = job.get('title', '').lower()
        description = job.get('description', '').lower()
        
        pref_level = self.preferences['experience_level'].lower()
        
        # Experience level keywords
        if pref_level == 'junior':
            if any(term in f"{title} {description}" for term in ['junior', 'entry', 'graduate']):
                return 100.0
            if any(term in f"{title} {description}" for term in ['senior', 'lead', 'principal']):
                return 20.0
        
        elif pref_level == 'mid':
            if any(term in f"{title} {description}" for term in ['mid', 'intermediate']):
                return 100.0
            if 'senior' in f"{title} {description}" and 'lead' not in f"{title} {description}":
                return 80.0
            if 'junior' in f"{title} {description}":
                return 60.0
        
        elif pref_level == 'senior':
            if any(term in f"{title} {description}" for term in ['senior', 'lead', 'principal', 'staff']):
                return 100.0
            if 'mid' in f"{title} {description}":
                return 60.0
        
        return 50.0  # Neutral if can't determine
    
    def _score_recency(self, job: Dict) -> float:
        """Score based on posting date (0-100)"""
        posted_date = job.get('posted_date', '')
        
        if not posted_date:
            return 50.0  # Neutral if no date
        
        try:
            if isinstance(posted_date, str):
                # Try to parse date
                posted = datetime.fromisoformat(posted_date.replace('Z', '+00:00'))
            else:
                posted = posted_date
            
            days_old = (datetime.now() - posted.replace(tzinfo=None)).days
            
            recency_threshold = self.preferences['recency_days']
            
            if days_old <= 7:
                return 100.0
            elif days_old <= 14:
                return 80.0
            elif days_old <= recency_threshold:
                return 60.0
            else:
                return 30.0
        
        except:
            return 50.0
    
    def _get_recommendation(self, score: float) -> str:
        """Get recommendation based on score"""
        if score >= 80:
            return "🔥 HIGH PRIORITY - Excellent match!"
        elif score >= 70:
            return "✅ GOOD MATCH - Worth applying"
        elif score >= 60:
            return "⚠️  MODERATE MATCH - Consider carefully"
        else:
            return "❌ LOW MATCH - Skip or lower priority"
    
    def score_jobs_batch(self, jobs: List[Dict]) -> List[Dict]:
        """Score multiple jobs"""
        scored_jobs = []
        
        for job in jobs:
            score_result = self.score_job(job)
            job_with_score = {**job, **score_result}
            scored_jobs.append(job_with_score)
        
        # Sort by score descending
        scored_jobs.sort(key=lambda x: x['total_score'], reverse=True)
        
        return scored_jobs


def main():
    """Main scoring function for workflow integration"""
    import argparse
    
    parser = argparse.ArgumentParser(description='Score job postings')
    parser.add_argument('--cv', '--input', dest='cv', required=True, help='Path to CV file (PDF, DOCX, or TXT)')
    parser.add_argument('--jobs', required=True, help='Path to jobs JSON file')
    parser.add_argument('--output', required=True, help='Output path for scored jobs')
    parser.add_argument('--preferences', help='Optional preferences JSON file')
    parser.add_argument('--min-score', type=float, default=60, help='Minimum score to keep')
    
    args = parser.parse_args()
    
    print("=" * 70)
    print("🎯 JOB SCORING SYSTEM")
    print("=" * 70)
    
    # Validate CV file exists
    cv_path = Path(args.cv)
    if not cv_path.exists():
        print(f"❌ CV file not found: {args.cv}")
        sys.exit(1)
    
    print(f"📄 CV File: {args.cv} ({cv_path.suffix})")
    
    # Initialize scorer
    try:
        scorer = JobScorer(args.cv, args.preferences)
    except Exception as e:
        print(f"❌ Failed to initialize scorer: {e}")
        sys.exit(1)
    
    # Load jobs
    jobs_file = Path(args.jobs)
    if not jobs_file.exists():
        print(f"❌ Jobs file not found: {args.jobs}")
        sys.exit(1)
    
    with open(jobs_file, 'r') as f:
        jobs = json.load(f)
    
    if isinstance(jobs, dict) and 'jobs' in jobs:
        jobs = jobs['jobs']
    
    print(f"📊 Loaded {len(jobs)} jobs to score")
    
    # Score jobs
    scored_jobs = scorer.score_jobs_batch(jobs)
    
    # Filter by minimum score
    filtered_jobs = [j for j in scored_jobs if j['total_score'] >= args.min_score]
    
    print(f"\n✅ Scored {len(scored_jobs)} jobs")
    print(f"🔍 {len(filtered_jobs)} jobs meet minimum score threshold ({args.min_score})")
    
    # Show top 5
    if filtered_jobs:
        print("\n🏆 TOP 5 MATCHES:")
        for i, job in enumerate(filtered_jobs[:5], 1):
            print(f"{i}. [{job['total_score']:.1f}] {job['job_title']} @ {job['company']}")
            print(f"   {job['recommendation']}")
    else:
        print("\n⚠️  No jobs met the minimum score threshold")
    
    # Save results
    output_file = Path(args.output)
    output_file.parent.mkdir(parents=True, exist_ok=True)
    
    with open(output_file, 'w') as f:
        json.dump({
            'scored_at': datetime.now().isoformat(),
            'total_jobs': len(jobs),
            'scored_jobs': len(scored_jobs),
            'filtered_jobs': len(filtered_jobs),
            'min_score': args.min_score,
            'top_matches': filtered_jobs[:10],
            'jobs': filtered_jobs
        }, f, indent=2)
    
    print(f"\n💾 Results saved to: {output_file}")
    print("=" * 70)


if __name__ == '__main__':
    main()
