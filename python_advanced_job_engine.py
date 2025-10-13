#!/usr/bin/env python3
"""
Advanced Job Engine - Complete Job Search & Skill Development System
Analyzes CV vs Job, Creates Learning Plans, Tracks Progress, Generates Applications
"""

import json
import re
from datetime import datetime, timedelta
from pathlib import Path
from typing import Dict, List, Tuple, Optional
from collections import defaultdict
import hashlib
import io

# PDF and DOCX reading libraries
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


    
class AdvancedJobEngine:
    def __init__(self, data_dir: str = "job_search_data"):
        self.data_dir = Path(data_dir)
        self.data_dir.mkdir(exist_ok=True)
        
        # File paths
        self.skillset_file = self.data_dir / "master_skillset.json"
        self.progress_file = self.data_dir / "learning_progress.json"
        self.jobs_file = self.data_dir / "analyzed_jobs.json"
        self.tests_file = self.data_dir / "skill_tests.json"
        self.sprints_file = self.data_dir / "sprint_history.json"
        self.state_file = self.data_dir / "workflow_state.json"
        
        # Load data
        self.master_skillset = self._load_json(self.skillset_file, self._init_skillset())
        self.learning_progress = self._load_json(self.progress_file, [])
        self.analyzed_jobs = self._load_json(self.jobs_file, [])
        self.skill_tests = self._load_json(self.tests_file, [])
        self.sprint_history = self._load_json(self.sprints_file, [])
        self.state = self._load_json(self.state_file, self._init_state())

        
        # Scoring weights
        self.WEIGHTS = {
            "required_skills": 0.35,
            "preferred_skills": 0.15,
            "experience": 0.20,
            "education": 0.10,
            "certifications": 0.05,
            "keywords": 0.15
        }
        
        # Learning resources database
        self.LEARNING_RESOURCES = self._init_learning_resources()
        
        # Test difficulty levels
        self.TEST_LEVELS = {
            "beginner": {"questions": 10, "pass_score": 60},
            "intermediate": {"questions": 15, "pass_score": 70},
            "advanced": {"questions": 20, "pass_score": 80}
        }

        # Quality gates (for reverse workflow)
        self.QUALITY_GATES = {
            "foundation": {"score": 65, "projects": 2, "tests_passed": "beginner"},
            "competency": {"score": 80, "projects": 4, "tests_passed": "intermediate"},
            "mastery": {"score": 90, "projects": 5, "tests_passed": "advanced"},
            "application_ready": {"score": 90, "brand": True, "network": True}
        }

        
    def _load_json(self, filepath: Path, default):
        if filepath.exists():
            with open(filepath, 'r', encoding='utf-8') as f:
                return json.load(f)
        return default
    
    def _save_json(self, filepath: Path, data):
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)


       def read_document(self, file_path: str) -> str:
        """
        Read document from file (supports .txt, .pdf, .docx)
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Extracted text content
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Determine file type and extract text
        ext = path.suffix.lower()
        
        if ext == '.txt':
            return self._read_txt(path)
        elif ext == '.pdf':
            return self._read_pdf(path)
        elif ext in ['.docx', '.doc']:
            return self._read_docx(path)
        else:
            raise ValueError(f"Unsupported file format: {ext}. Supported: .txt, .pdf, .docx")
    
    def _read_txt(self, path: Path) -> str:
        """Read plain text file"""
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    
    def _read_pdf(self, path: Path) -> str:
        """Read PDF file"""
        if not PDF_AVAILABLE:
            raise ImportError(
                "PyPDF2 is required to read PDF files. "
                "Install it with: pip install PyPDF2"
            )
        
        text = []
        try:
            with open(path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page in pdf_reader.pages:
                    text.append(page.extract_text())
            return '\n'.join(text)
        except Exception as e:
            raise Exception(f"Error reading PDF: {str(e)}")
    
    def _read_docx(self, path: Path) -> str:
        """Read DOCX file"""
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required to read DOCX files. "
                "Install it with: pip install python-docx"
            )
        
        try:
            doc = Document(path)
            text = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text.append(cell.text)
            
            return '\n'.join(text)
        except Exception as e:
            raise Exception(f"Error reading DOCX: {str(e)}")
            
    def _init_skillset(self) -> Dict:
        """Initialize master skillset structure"""
        return {
            "technical_skills": {
                "programming": [],
                "frameworks": [],
                "tools": [],
                "databases": [],
                "cloud": []
            },
            "soft_skills": [],
            "certifications": [],
            "languages": [],
            "domains": [],
            "last_updated": datetime.now().isoformat()
        }
    
    def _init_learning_resources(self) -> Dict:
        """Initialize comprehensive learning resources database"""
        return {
            "python": {
                "study": ["Python.org Documentation", "Real Python Tutorials", "Python Crash Course Book"],
                "practice": ["LeetCode Python", "HackerRank Python", "Codewars Python"],
                "courses": ["Python for Everybody (Coursera)", "Complete Python Bootcamp (Udemy)", "MIT 6.0001"]
            },
            "machine learning": {
                "study": ["Hands-On ML Book", "ML Crash Course (Google)", "Scikit-learn Docs"],
                "practice": ["Kaggle Competitions", "ML Project Ideas", "Scikit-learn Examples"],
                "courses": ["Machine Learning (Coursera)", "Fast.ai", "Deep Learning Specialization"]
            },
            "pytorch": {
                "study": ["PyTorch Documentation", "PyTorch Tutorials", "Deep Learning with PyTorch Book"],
                "practice": ["PyTorch Examples Repo", "Papers with Code", "Kaggle PyTorch Kernels"],
                "courses": ["Deep Learning with PyTorch (Udacity)", "PyTorch for Deep Learning (Udemy)"]
            },
            "aws": {
                "study": ["AWS Documentation", "AWS Well-Architected", "AWS Whitepapers"],
                "practice": ["AWS Free Tier Projects", "AWS Hands-On Labs", "LocalStack"],
                "courses": ["AWS Solutions Architect (Udemy)", "A Cloud Guru AWS", "AWS Training Portal"]
            },
            "docker": {
                "study": ["Docker Documentation", "Docker Deep Dive Book", "Docker Getting Started"],
                "practice": ["Docker Labs", "Dockerize Projects", "Docker Compose Examples"],
                "courses": ["Docker Mastery (Udemy)", "Docker and Kubernetes (Udemy)"]
            },
            "kubernetes": {
                "study": ["Kubernetes Docs", "Kubernetes Up & Running Book", "K8s Patterns"],
                "practice": ["Minikube Labs", "K8s the Hard Way", "Kubernetes Examples"],
                "courses": ["CKA Certification Prep", "Kubernetes for Developers", "K8s Mastery"]
            },
            "sql": {
                "study": ["W3Schools SQL", "PostgreSQL Tutorial", "SQL Performance Explained"],
                "practice": ["SQLZoo", "LeetCode SQL", "HackerRank SQL"],
                "courses": ["Complete SQL Bootcamp", "SQL for Data Science", "Advanced SQL (Mode)"]
            },
            "javascript": {
                "study": ["MDN Web Docs", "JavaScript.info", "Eloquent JavaScript"],
                "practice": ["JavaScript30", "FreeCodeCamp JS", "Codewars JS"],
                "courses": ["JavaScript Complete Course", "Modern JavaScript (Udemy)", "JS Algorithms"]
            },
            "nlp": {
                "study": ["NLP with Python Book", "Hugging Face Docs", "Speech & Language Processing"],
                "practice": ["Kaggle NLP Challenges", "NLP Projects", "Hugging Face Tasks"],
                "courses": ["NLP Specialization (Coursera)", "Fast.ai NLP", "Advanced NLP (Udemy)"]
            },
            "default": {
                "study": ["Official Documentation", "Technical Blogs", "Research Papers"],
                "practice": ["GitHub Projects", "Coding Challenges", "Real-world Applications"],
                "courses": ["Coursera Courses", "Udemy Courses", "YouTube Tutorials"]
            }
        }
    
    def parse_cv(self, cv_text: str) -> Dict:
        """Extract structured information from CV"""
        return {
            "raw_text": cv_text,
            "skills": self._extract_all_skills(cv_text),
            "experience_years": self._extract_experience(cv_text),
            "education": self._extract_education(cv_text),
            "certifications": self._extract_certifications(cv_text),
            "projects": self._extract_projects(cv_text),
            "soft_skills": self._extract_soft_skills(cv_text)
        }
    
    def parse_job(self, job_text: str, title: str = "", company: str = "") -> Dict:
        """Extract structured requirements from job description"""
        return {
            "id": hashlib.md5((title + company + job_text).encode()).hexdigest()[:12],
            "title": title or self._extract_title(job_text),
            "company": company or self._extract_company(job_text),
            "raw_text": job_text,
            "required_skills": self._extract_required_skills(job_text),
            "preferred_skills": self._extract_preferred_skills(job_text),
            "required_experience": self._extract_required_experience(job_text),
            "education_required": self._extract_education_req(job_text),
            "certifications_required": self._extract_cert_req(job_text),
            "keywords": self._extract_keywords(job_text),
            "responsibilities": self._extract_responsibilities(job_text),
            "analyzed_date": datetime.now().isoformat()
        }



    def analyze_from_files(self, cv_file: str, job_file: str, 
                          job_title: str = "", company: str = "") -> Dict:
        """
        Analyze job match from document files
        
        Args:
            cv_file: Path to CV file (.txt, .pdf, .docx)
            job_file: Path to job description file (.txt, .pdf, .docx)
            job_title: Optional job title
            company: Optional company name
            
        Returns:
            Complete analysis dictionary
        """
        print(f"📄 Reading CV from: {cv_file}")
        cv_text = self.read_document(cv_file)
        print(f"   ✅ Extracted {len(cv_text)} characters")
        
        print(f"📄 Reading job description from: {job_file}")
        job_text = self.read_document(job_file)
        print(f"   ✅ Extracted {len(job_text)} characters")
        
        print(f"\n📊 Analyzing job match...")
        return self.analyze_job_complete(cv_text, job_text, job_title, company)
    
    def analyze_job_complete(self, cv_text: str, job_text: str, 
                            job_title: str = "", company: str = "") -> Dict:
        """
        STEP 1 & 2: Complete job analysis with scoring
        """
        cv_data = self.parse_cv(cv_text)
        job_data = self.parse_job(job_text, job_title, company)
        
        # Calculate match score
        score_result = self._calculate_score(cv_data, job_data)
        
        # Identify gaps
        gaps = self._identify_gaps(cv_data, job_data)
        
        # Generate recommendations
        recommendations = self._generate_recommendations(gaps, score_result)
        
        analysis = {
            "job_id": job_data["id"],
            "job_info": {
                "title": job_data["title"],
                "company": job_data["company"]
            },
            "score": score_result,
            "gaps": gaps,
            "recommendations": recommendations,
            "cv_snapshot": cv_data,
            "job_snapshot": job_data,
            "analysis_date": datetime.now().isoformat()
        }
        
        # Save analysis
        self.analyzed_jobs.append(analysis)
        self._save_json(self.jobs_file, self.analyzed_jobs)
        
        return analysis


    def create_learning_plan(self, analysis: Dict, mode: str = "standard") -> Dict:
        """
        STEP 3: Create detailed learning plan with three levels
    
        Args:
        analysis: Job analysis dictionary
        mode: "standard" (12 weeks, quick) or "reverse" (16-24 weeks, mastery-focused)
    
        Returns:
        Comprehensive learning plan
        """
        gaps = analysis["gaps"]
    
        # Set duration based on mode
        duration = "12 weeks" if mode == "standard" else "16-24 weeks"
    
        plan = {
            "plan_id": f"LP_{analysis['job_id']}_{datetime.now().strftime('%Y%m%d')}",
            "mode": mode,  # ADD THIS LINE
            "job_reference": {
                "title": analysis["job_info"]["title"],
                "company": analysis["job_info"]["company"]
            },
            "created_date": datetime.now().isoformat(),
            "estimated_duration": duration,  # MODIFY THIS LINE
            "levels": {
                "study": [],
                "practice": [],
                "courses": []
            },
            "weekly_schedule": [],
            "milestones": []
        }
        
        # Categorize missing skills
        missing_required = gaps["missing_required_skills"]
        missing_preferred = gaps["missing_preferred_skills"]
        
        # Level A: TO STUDY (completely new skills)
        for skill in missing_required[:5]:  # Top 5 critical skills
            resources = self._get_resources(skill)
            plan["levels"]["study"].append({
                "skill": skill,
                "priority": "CRITICAL",
                "category": "required",
                "resources": resources["study"],
                "estimated_time": "2-3 weeks",
                "prerequisites": self._get_prerequisites(skill),
                "learning_path": self._generate_learning_path(skill)
            })
        
        # Level B: TO PRACTICE (have some knowledge, need more practice)
        for skill in missing_preferred[:3]:
            resources = self._get_resources(skill)
            plan["levels"]["practice"].append({
                "skill": skill,
                "priority": "MEDIUM",
                "category": "preferred",
                "practice_activities": resources["practice"],
                "estimated_time": "1-2 weeks",
                "practice_goals": self._generate_practice_goals(skill)
            })
        
        # Level C: COURSES (structured learning)
        all_missing = list(set(missing_required + missing_preferred))
        for skill in all_missing[:7]:  # Top 7 skills
            resources = self._get_resources(skill)
            plan["levels"]["courses"].append({
                "skill": skill,
                "recommended_courses": resources["courses"],
                "estimated_duration": "4-8 weeks per course",
                "certification_available": self._check_certification(skill),
                "priority": "HIGH" if skill in missing_required else "MEDIUM"
            })
        
        # Generate weekly schedule
        plan["weekly_schedule"] = self._generate_weekly_schedule(plan)
        
        # Set milestones
        plan["milestones"] = self._generate_milestones(plan)
        
        # Save plan
        self.learning_progress.append(plan)
        self._save_json(self.progress_file, self.learning_progress)
        
        return plan

    # ====================
    # ADD NEW METHOD: _init_state
    # ====================

    def _init_state(self) -> Dict:
        """Initialize workflow state"""
        return {
            "mode": None,  # "standard" or "reverse"
            "current_stage": "baseline",  # baseline, skill_building, mastery, positioning, ready
            "baseline_score": 0,
            "current_score": 0,
            "target_score": 90,
            "current_sprint": 0,
            "skills_mastered": [],
            "projects_completed": [],
            "tests_passed": {},  # Track test levels passed per skill
            "quality_gates_passed": [],
            "brand_ready": False,
            "network_ready": False,
            "application_ready": False,
            "started_date": None,
            "estimated_completion": None
        }


    # ====================
    # ADD NEW METHOD: start_sprint
    # ====================

    def start_sprint(self, skills: List[str], project_goal: str) -> Dict:
        """
        REVERSE WORKFLOW: Start a 2-week sprint for skill building
        
        Args:
        skills: List of skills to focus on during sprint
        project_goal: Description of project to build
    
        Returns:
        Sprint configuration dictionary
        """
        sprint_num = self.state["current_sprint"] + 1
    
        sprint = {
            "sprint_number": sprint_num,
            "start_date": datetime.now().isoformat(),
            "end_date": (datetime.now() + timedelta(days=14)).isoformat(),
            "skills_targeted": skills,
            "project_goal": project_goal,
            "daily_logs": [],
            "completed": False
        }
    
        self.state["current_sprint"] = sprint_num
        self.state["mode"] = "reverse"
        if not self.state["started_date"]:
            self.state["started_date"] = datetime.now().isoformat()
        self._save_json(self.state_file, self.state)
    
        self.sprint_history.append(sprint)
        self._save_json(self.sprints_file, self.sprint_history)
    
        print(f"\n🏃 Sprint {sprint_num} Started")
        print(f"   Skills: {', '.join(skills)}")
        print(f"   Project: {project_goal}")
        print(f"   Duration: 14 days")
        print(f"   End Date: {sprint['end_date'][:10]}")
    
        return sprint


    # ====================
    # ADD NEW METHOD: log_daily
    # ====================

    def log_daily(self, hours: float, concepts: List[str], notes: str = ""):
        """
        Log daily learning activity for current sprint
    
        Args:
        hours: Hours spent learning today
        concepts: List of concepts/topics covered
        notes: Optional notes about the day's progress
        """
        if not self.sprint_history:
            print("⚠️ No active sprint. Start a sprint first with start_sprint()")
            return
    
        current_sprint = self.sprint_history[-1]
    
        if current_sprint.get("completed", False):
            print("⚠️ Current sprint is completed. Start a new sprint.")
            return
    
        log_entry = {
            "date": datetime.now().isoformat(),
            "day_number": len(current_sprint["daily_logs"]) + 1,
            "hours": hours,
            "concepts": concepts,
            "notes": notes
        }
    
        current_sprint["daily_logs"].append(log_entry)
        self._save_json(self.sprints_file, self.sprint_history)
    
        total_hours = sum(log["hours"] for log in current_sprint["daily_logs"])
    
        print(f"✅ Day {log_entry['day_number']} logged ({hours}h)")
        print(f"   Total Sprint Hours: {total_hours}h")
        print(f"   Concepts: {', '.join(concepts)}")


        # ====================
        # ADD NEW METHOD: end_sprint
        # ====================

        def end_sprint(self, project_url: str, test_scores: Dict[str, float]) -> Dict:
            """
            Complete current sprint and assess progress
    
            Args:
            project_url: URL to completed project (GitHub, portfolio, etc.)
            test_scores: Dictionary of {skill: score} from skill tests
    
            Returns:
            Sprint completion summary
            """
            if not self.sprint_history:
                print("⚠️ No active sprint")
                return {}
    
            current_sprint = self.sprint_history[-1]
    
            if current_sprint.get("completed", False):
                print("⚠️ Sprint already completed")
                return {}
    
            sprint_num = current_sprint["sprint_number"]
    
            print(f"\n🏁 Sprint {sprint_num} Assessment")
            print("=" * 60)
    
            # Record results
            current_sprint["completed"] = True
            current_sprint["completion_date"] = datetime.now().isoformat()
            current_sprint["project_url"] = project_url
            current_sprint["test_scores"] = test_scores
    
            total_hours = sum(log["hours"] for log in current_sprint["daily_logs"])
            current_sprint["total_hours"] = total_hours
    
            # Update mastered skills
            newly_mastered = []
            for skill, score in test_scores.items():
                if score >= 60 and skill not in self.state["skills_mastered"]:
                    self.state["skills_mastered"].append(skill)
                    newly_mastered.append(skill)
                    print(f"   ✅ Mastered: {skill} (Score: {score}%)")
        
                # Track test level passed
                if skill not in self.state["tests_passed"]:
                    self.state["tests_passed"][skill] = []
        
                if score >= 80:
                    level = "advanced"
                elif score >= 70:
                    level = "intermediate"
                elif score >= 60:
                    level = "beginner"
                else:
                    level = "needs_improvement"
        
                if level not in self.state["tests_passed"][skill]:
                    self.state["tests_passed"][skill].append(level)
    
            # Add project
            project_entry = {
                "sprint": sprint_num,
                "goal": current_sprint["project_goal"],
                "url": project_url,
                "skills": current_sprint["skills_targeted"],
                "completion_date": datetime.now().isoformat()
            }
            self.state["projects_completed"].append(project_entry)
    
            self._save_json(self.sprints_file, self.sprint_history)
            self._save_json(self.state_file, self.state)
    
            # Check quality gates
            gates_passed = self.check_quality_gates()
    
            # Summary
            summary = {
                "sprint": sprint_num,
                "total_hours": total_hours,
                "skills_mastered": newly_mastered,
                "project_completed": True,
                "test_scores": test_scores,
                "quality_gates_passed": [g for g, passed in gates_passed.items() if passed]
            }
    
            print(f"\n📊 Sprint {sprint_num} Summary:")
            print(f"   Total Hours: {total_hours}h")
            print(f"   Skills Mastered: {len(self.state['skills_mastered'])}")
            print(f"   Projects Completed: {len(self.state['projects_completed'])}")
            print(f"   Quality Gates: {', '.join(summary['quality_gates_passed']) or 'None yet'}")
    
            return summary


        # ====================
        # ADD NEW METHOD: check_quality_gates
        # ====================

        def check_quality_gates(self) -> Dict[str, bool]:
            """
            Check which quality gates have been passed
            
            Returns:
            Dictionary of {gate_name: passed_status}
            """
            score = self.state["current_score"]
            projects = len(self.state["projects_completed"])
    
            gates_status = {}
            newly_passed = []
    
            for gate_name, requirements in self.QUALITY_GATES.items():
                # Check score and projects
                score_met = score >= requirements["score"]
                projects_met = projects >= requirements["projects"]
        
                # Check additional requirements
                if "brand" in requirements:
                    additional_met = self.state.get("brand_ready", False)
                elif "tests_passed" in requirements:
                    level = requirements["tests_passed"]
                    # Check if at least one skill has passed this level
                    additional_met = any(
                        level in levels 
                        for levels in self.state["tests_passed"].values()
                    )
                else:
                    additional_met = True
        
                passed = score_met and projects_met and additional_met
                gates_status[gate_name] = passed
        
                # Track newly passed gates
                if passed and gate_name not in self.state["quality_gates_passed"]:
                    self.state["quality_gates_passed"].append(gate_name)
                    newly_passed.append(gate_name)
                    print(f"   🏆 Quality Gate Passed: {gate_name.upper()}")
            
                    # Update stage
                    if gate_name == "foundation":
                        self.state["current_stage"] = "skill_building"
                    elif gate_name == "competency":
                        self.state["current_stage"] = "mastery"
                    elif gate_name == "mastery":
                        self.state["current_stage"] = "positioning"
                    elif gate_name == "application_ready":
                        self.state["current_stage"] = "ready"
    
            if newly_passed:
                self._save_json(self.state_file, self.state)
    
            return gates_status


        # ====================
        # ADD NEW METHOD: stage_positioning
        # ====================

        def stage_positioning(self) -> Dict:
            """
            STAGE 5: Build professional brand and visibility (before applying)
            
            Returns:
            Comprehensive checklist for market positioning
            """
            print("\n" + "="*80)
            print("STAGE 5: MARKET POSITIONING & VISIBILITY")
            print("="*80)
            
            if self.state["current_score"] < 85:
                print(f"\n⚠️ Current Score: {self.state['current_score']}%")
                print(f"   Recommended: Reach 85%+ before positioning stage")
                print(f"   Continue skill building and project work")
                return {}
    
            checklist = {
                "linkedin_profile": {
                    "priority": "CRITICAL",
                    "tasks": [
                        "Add all new skills to Skills section",
                        "Showcase 5+ portfolio projects in Featured section",
                        "Write detailed experience descriptions with metrics",
                        "Professional photo and compelling headline",
                        "Get 3+ recommendations from colleagues/mentors",
                        "Join relevant groups and engage weekly",
                        "Share technical content regularly"
                    ]
                },
                "github_portfolio": {
                    "priority": "CRITICAL",
                    "tasks": [
                        "Organize repositories professionally with topics/tags",
                        "Comprehensive READMEs for all major projects",
                        "Pin best 6 projects to profile",
                        "Active contribution graph (green squares daily)",
                        "Profile README with portfolio showcase and bio",
                        "Clean up code, add comments and documentation",
                        "Include live demos or screenshots"
                    ]
                },
                "technical_content": {
                    "priority": "HIGH",
                    "tasks": [
                        "Write 3-5 technical blog posts (Medium, Dev.to)",
                        "Create tutorial or guide in your specialty",
                        "Document your learning journey publicly",
                        "Share insights on LinkedIn 2-3x per week",
                        "Answer questions on Stack Overflow",
                        "Contribute to technical discussions"
                    ]
                },
                "personal_website": {
                    "priority": "MEDIUM",
                    "tasks": [
                        "Portfolio showcase page with project cards",
                        "About/story section telling your journey",
                        "Contact information and social links",
                        "Resume download (PDF)",
                        "Blog or articles section (optional)",
                        "Mobile-responsive design",
                        "SEO optimization"
                    ]
                },
                "networking": {
                    "priority": "HIGH",
                    "tasks": [
                        "Connect with 20+ relevant professionals on LinkedIn",
                        "Join 3+ industry communities/Slack groups",
                        "Attend 2+ virtual meetups or webinars monthly",
                        "Engage with target company employees",
                        "Request 5+ informational interviews",
                        "Follow up and maintain relationships",
                        "Offer value before asking for help"
                    ]
                }
            }
            
            print("\n📋 Professional Branding Checklist:\n")
            for area, info in checklist.items():
                print(f"📌 {area.replace('_', ' ').title()} [{info['priority']}]:")
                for task in info["tasks"]:
                    print(f"   ☐ {task}")
                print()
    
            print("=" * 80)
            print("TRACKING YOUR POSITIONING PROGRESS:")
            print("=" * 80)
            print("""
            Mark tasks as complete, then update state:

            # When LinkedIn is polished:
            engine.state['linkedin_ready'] = True
  
            # When GitHub portfolio is professional:
            engine.state['github_ready'] = True
  
            # When overall brand is ready:
            engine.state['brand_ready'] = True
  
            # When network is established:
            engine.state['network_ready'] = True
  
            # Save state:
            engine._save_json(engine.state_file, engine.state)
  
            # Check if application ready:
            engine.check_quality_gates()
            """)
    
            print("\n💡 PRO TIPS:")
            print("   • Quality over quantity - 5 great projects > 20 mediocre ones")
            print("   • Consistency matters - regular activity shows dedication")
            print("   • Engagement beats broadcasting - comment, discuss, help others")
            print("   • Tell your story - people connect with journeys, not just skills")
            print("   • Be authentic - genuine passion shines through")
    
            return checklist


        # ====================
        # ADD NEW METHOD: display_progress_dashboard
        # ====================

        def display_progress_dashboard(self):
            """Display comprehensive progress dashboard"""
            print("\n" + "="*80)
            print("PROGRESS DASHBOARD")
            print("="*80)
    
            state = self.state
    
            print(f"\n📊 CURRENT STATUS:")
            print(f"   Mode: {state['mode'] or 'Not started'}")
            print(f"   Stage: {state['current_stage']}")
            print(f"   Baseline Score: {state['baseline_score']}%")
            print(f"   Current Score: {state['current_score']}%")
            print(f"   Target Score: {state['target_score']}%")
    
            if state['baseline_score'] > 0:
                improvement = state['current_score'] - state['baseline_score']
                print(f"   Improvement: +{improvement}%")
    
            print(f"\n🏃 SPRINT PROGRESS:")
            print(f"   Total Sprints: {state['current_sprint']}")
            print(f"   Skills Mastered: {len(state['skills_mastered'])}")
            if state['skills_mastered']:
                print(f"      → {', '.join(state['skills_mastered'][:10])}")
    
            print(f"\n🗂️ PROJECTS COMPLETED: {len(state['projects_completed'])}")
            for i, project in enumerate(state['projects_completed'][-5:], 1):
                print(f"   {i}. Sprint {project['sprint']}: {project['goal']}")
                print(f"      Skills: {', '.join(project['skills'])}")
                print(f"      URL: {project['url']}")
                
            print(f"\n📝 TESTS PASSED:")
            if state['tests_passed']:
                for skill, levels in state['tests_passed'].items():
                    print(f"   {skill}: {', '.join(levels)}")
            else:
                print("   No tests passed yet")
    
            print(f"\n🏆 QUALITY GATES:")
            for gate in self.QUALITY_GATES.keys():
                status = "✅" if gate in state['quality_gates_passed'] else "⏳"
                print(f"   {status} {gate.replace('_', ' ').title()}")
    
            print(f"\n🎯 READINESS FLAGS:")
            print(f"   Brand Ready: {'✅' if state.get('brand_ready') else '⏳'}")
            print(f"   Network Ready: {'✅' if state.get('network_ready') else '⏳'}")
            print(f"   Application Ready: {'✅' if state.get('application_ready') else '⏳'}")
    
            if state['started_date']:
                start = datetime.fromisoformat(state['started_date'])
                days_elapsed = (datetime.now() - start).days
                print(f"\n⏱️ TIME TRACKING:")
                print(f"   Started: {start.strftime('%Y-%m-%d')}")
                print(f"   Days Elapsed: {days_elapsed}")
        
                if self.sprint_history:
                    total_hours = sum(
                        sprint.get('total_hours', 0) 
                        for sprint in self.sprint_history 
                        if sprint.get('completed')
                    )
                    print(f"   Total Learning Hours: {total_hours}h")
    
            print("\n" + "="*80)

          
    def create_improvement_strategy(self, analysis: Dict, learning_plan: Dict) -> Dict:
        """
        STEP 4: Create detailed plan to overcome gaps
        """
        strategy = {
            "strategy_id": f"STR_{analysis['job_id']}",
            "created_date": datetime.now().isoformat(),
            "target_job": analysis["job_info"],
            "current_score": analysis["score"]["total_score"],
            "target_score": 80,  # Aim for 80%+
            "timeline": "12 weeks",
            "phases": []
        }
        
        # Phase 1: Foundation (Weeks 1-4)
        strategy["phases"].append({
            "phase": 1,
            "name": "Foundation Building",
            "duration": "4 weeks",
            "focus": "Critical required skills",
            "objectives": [
                f"Master {skill}" for skill in analysis["gaps"]["missing_required_skills"][:3]
            ],
            "activities": [
                "Complete foundational courses",
                "Study documentation and tutorials",
                "Build 2-3 small projects"
            ],
            "success_criteria": [
                "Pass beginner level tests",
                "Complete 3 hands-on projects",
                "Demonstrate basic proficiency"
            ],
            "expected_score_increase": 15
        })
        
        # Phase 2: Skill Development (Weeks 5-8)
        strategy["phases"].append({
            "phase": 2,
            "name": "Skill Development",
            "duration": "4 weeks",
            "focus": "Advanced skills and practice",
            "objectives": [
                "Deepen understanding of core technologies",
                "Build portfolio projects",
                "Gain practical experience"
            ],
            "activities": [
                "Complete intermediate courses",
                "Contribute to open source",
                "Build 2-3 medium complexity projects",
                "Practice coding challenges daily"
            ],
            "success_criteria": [
                "Pass intermediate level tests",
                "Complete 3 portfolio projects",
                "Demonstrate intermediate proficiency"
            ],
            "expected_score_increase": 15
        })
        
        # Phase 3: Mastery & Application (Weeks 9-12)
        strategy["phases"].append({
            "phase": 3,
            "name": "Mastery & Job Application",
            "duration": "4 weeks",
            "focus": "Advanced topics and job preparation",
            "objectives": [
                "Master all required skills",
                "Build comprehensive portfolio",
                "Prepare for interviews"
            ],
            "activities": [
                "Complete advanced courses",
                "Build 1-2 advanced projects",
                "Practice system design",
                "Mock interviews",
                "Update CV and LinkedIn"
            ],
            "success_criteria": [
                "Pass advanced level tests",
                "Have 5+ strong portfolio projects",
                "Ready for interviews"
            ],
            "expected_score_increase": 10
        })
        
        # Action items
        strategy["action_items"] = self._generate_action_items(analysis, learning_plan)
        
        # Progress tracking
        strategy["progress_tracking"] = {
            "weekly_check_ins": True,
            "skill_assessments": "Every 2 weeks",
            "portfolio_reviews": "Monthly",
            "mock_interviews": "Week 10 & 12"
        }
        
        return strategy
    
    def generate_skill_tests(self, skills: List[str]) -> Dict:
        """
        STEP 5: Generate three-level tests to measure progress
        - Beginner
        - Intermediate  
        - Advanced
        """
        tests = {
            "test_suite_id": f"TEST_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
            "created_date": datetime.now().isoformat(),
            "skills_covered": skills,
            "levels": {}
        }
        
        for skill in skills:
            tests["levels"][skill] = {
                "beginner": self._generate_test(skill, "beginner"),
                "intermediate": self._generate_test(skill, "intermediate"),
                "advanced": self._generate_test(skill, "advanced")
            }
        
        # Save tests
        self.skill_tests.append(tests)
        self._save_json(self.tests_file, self.skill_tests)
        
        return tests
    
    def _generate_test(self, skill: str, level: str) -> Dict:
        """Generate test for specific skill and level"""
        config = self.TEST_LEVELS[level]
        
        test = {
            "skill": skill,
            "level": level,
            "questions": config["questions"],
            "pass_score": config["pass_score"],
            "time_limit": config["questions"] * 5,  # 5 mins per question
            "question_types": []
        }
        
        # Generate question templates based on skill and level
        if level == "beginner":
            test["question_types"] = [
                f"Basic syntax and concepts of {skill}",
                f"Fundamental operations in {skill}",
                f"Common use cases for {skill}",
                f"Simple problem solving with {skill}",
                f"Understanding {skill} documentation"
            ]
            test["sample_questions"] = self._generate_beginner_questions(skill)
        
        elif level == "intermediate":
            test["question_types"] = [
                f"Advanced concepts in {skill}",
                f"Problem solving with {skill}",
                f"Best practices and patterns",
                f"Integration with other technologies",
                f"Performance optimization"
            ]
            test["sample_questions"] = self._generate_intermediate_questions(skill)
        
        else:  # advanced
            test["question_types"] = [
                f"Expert-level {skill} concepts",
                f"System design with {skill}",
                f"Complex problem solving",
                f"Architecture decisions",
                f"Production considerations"
            ]
            test["sample_questions"] = self._generate_advanced_questions(skill)
        
        test["assessment_criteria"] = {
            "technical_accuracy": 40,
            "problem_solving": 30,
            "best_practices": 20,
            "code_quality": 10
        }
        
        return test
    
    def _generate_beginner_questions(self, skill: str) -> List[Dict]:
        """Generate beginner level questions"""
        templates = {
            "python": [
                {"q": "What is the difference between a list and tuple in Python?", "type": "concept"},
                {"q": "Write a function to reverse a string", "type": "coding"},
                {"q": "Explain what a dictionary is and give an example", "type": "concept"},
                {"q": "How do you handle exceptions in Python?", "type": "concept"},
                {"q": "Write a program to find the sum of numbers in a list", "type": "coding"}
            ],
            "sql": [
                {"q": "What is the difference between WHERE and HAVING?", "type": "concept"},
                {"q": "Write a query to select all users with age > 25", "type": "coding"},
                {"q": "Explain what a JOIN is", "type": "concept"},
                {"q": "How do you sort results in ascending order?", "type": "concept"},
                {"q": "Write a query to count records in a table", "type": "coding"}
            ],
            "default": [
                {"q": f"What are the basic concepts of {skill}?", "type": "concept"},
                {"q": f"How do you get started with {skill}?", "type": "concept"},
                {"q": f"What are common use cases for {skill}?", "type": "concept"},
                {"q": f"Solve a simple problem using {skill}", "type": "coding"},
                {"q": f"Explain key terminology in {skill}", "type": "concept"}
            ]
        }
        return templates.get(skill.lower(), templates["default"])[:5]
    
    def _generate_intermediate_questions(self, skill: str) -> List[Dict]:
        """Generate intermediate level questions"""
        return [
            {"q": f"Design a solution for [problem] using {skill}", "type": "design"},
            {"q": f"What are the performance considerations when using {skill}?", "type": "concept"},
            {"q": f"How do you debug issues in {skill}?", "type": "practical"},
            {"q": f"Implement [complex feature] using {skill}", "type": "coding"},
            {"q": f"Compare {skill} with alternative solutions", "type": "analysis"}
        ]
    
    def _generate_advanced_questions(self, skill: str) -> List[Dict]:
        """Generate advanced level questions"""
        return [
            {"q": f"Design a scalable system using {skill}", "type": "system_design"},
            {"q": f"How would you optimize {skill} for production?", "type": "optimization"},
            {"q": f"Explain advanced architectural patterns with {skill}", "type": "architecture"},
            {"q": f"Solve a complex real-world problem with {skill}", "type": "coding"},
            {"q": f"What are the trade-offs when choosing {skill}?", "type": "analysis"}
        ]
    
    def update_skillset(self, new_skills: List[str], category: str = "technical") -> Dict:
        """
        STEP 6: Update master skillset with newly acquired skills
        """
        updated = []
        
        for skill in new_skills:
            skill_lower = skill.lower().strip()
            
            # Determine subcategory
            if category == "technical":
                if skill_lower in ["python", "java", "javascript", "c++", "go", "rust"]:
                    subcategory = "programming"
                elif skill_lower in ["pytorch", "tensorflow", "react", "django", "flask"]:
                    subcategory = "frameworks"
                elif skill_lower in ["docker", "kubernetes", "git", "jenkins"]:
                    subcategory = "tools"
                elif skill_lower in ["postgresql", "mongodb", "redis", "mysql"]:
                    subcategory = "databases"
                elif skill_lower in ["aws", "azure", "gcp"]:
                    subcategory = "cloud"
                else:
                    subcategory = "tools"
                
                if skill not in self.master_skillset["technical_skills"][subcategory]:
                    self.master_skillset["technical_skills"][subcategory].append(skill)
                    updated.append(skill)
            
            elif category == "certification":
                if skill not in self.master_skillset["certifications"]:
                    self.master_skillset["certifications"].append(skill)
                    updated.append(skill)
            
            elif category == "soft":
                if skill not in self.master_skillset["soft_skills"]:
                    self.master_skillset["soft_skills"].append(skill)
                    updated.append(skill)
        
        # Update timestamp
        self.master_skillset["last_updated"] = datetime.now().isoformat()
        
        # Save updated skillset
        self._save_json(self.skillset_file, self.master_skillset)
        
        return {
            "updated_skills": updated,
            "total_skills": self._count_total_skills(),
            "last_updated": self.master_skillset["last_updated"]
        }
    
    def generate_recruiter_letter(self, analysis: Dict, learning_plan: Dict = None) -> Dict:
        """
        STEP 7: Generate customized recruiter letters based on analysis
        """
        job_info = analysis["job_info"]
        score = analysis["score"]["total_score"]
        gaps = analysis["gaps"]
        
        letters = {
            "generated_date": datetime.now().isoformat(),
            "job_reference": job_info,
            "templates": {}
        }
        
        # Template 1: Strong Match Cover Letter (75%+)
        if score >= 75:
            letters["templates"]["cover_letter"] = self._generate_strong_match_letter(analysis)
        
        # Template 2: Growth Potential Letter (60-74%)
        elif score >= 60:
            letters["templates"]["cover_letter"] = self._generate_growth_letter(analysis)
        
        # Template 3: Future Interest Letter (<60%)
        else:
            letters["templates"]["cover_letter"] = self._generate_future_interest_letter(analysis, learning_plan)
        
        # Template 4: LinkedIn Connection Request
        letters["templates"]["linkedin_message"] = self._generate_linkedin_message(analysis)
        
        # Template 5: Follow-up Email
        letters["templates"]["followup_email"] = self._generate_followup_email(analysis)
        
        # Template 6: Networking Email
        letters["templates"]["networking_email"] = self._generate_networking_email(analysis)
        
        # Save letters
        filename = f"recruiter_letters_{analysis['job_id']}.json"
        self._save_json(self.data_dir / filename, letters)
        
        return letters
    
    def _generate_strong_match_letter(self, analysis: Dict) -> str:
        """Generate cover letter for strong matches"""
        job = analysis["job_info"]
        cv = analysis["cv_snapshot"]
        
        letter = f"""Dear Hiring Manager,

I am writing to express my strong interest in the {job['title']} position at {job['company']}. With {cv['experience_years']} years of relevant experience and a proven track record in the key areas you're seeking, I am confident I would be a valuable addition to your team.

KEY QUALIFICATIONS:

"""
        # Add matching skills
        matching_skills = []
        for category in cv['skills'].values():
            matching_skills.extend(category[:3])
        
        for skill in matching_skills[:5]:
            letter += f"✓ Proficient in {skill} with hands-on project experience\n"
        
        letter += f"""
RELEVANT EXPERIENCE:

I have successfully delivered projects involving:
"""
        
        for project in cv['projects'][:3]:
            letter += f"• {project}\n"
        
        letter += f"""
I am particularly excited about this opportunity because it aligns perfectly with my career goals and expertise. I am immediately available and eager to contribute to {job['company']}'s success.

I would welcome the opportunity to discuss how my skills and experience can benefit your team. Thank you for considering my application.

Best regards,
[Your Name]
[Your Phone]
[Your Email]
[Your LinkedIn]

---
ATTACHMENTS:
- Resume/CV
- Portfolio (if applicable)
- References (upon request)
"""
        return letter
    
    def _generate_growth_letter(self, analysis: Dict) -> str:
        """Generate cover letter emphasizing growth potential"""
        job = analysis["job_info"]
        cv = analysis["cv_snapshot"]
        gaps = analysis["gaps"]
        
        letter = f"""Dear Hiring Manager,

I am excited to apply for the {job['title']} position at {job['company']}. While I am actively developing expertise in some of the technologies you require, I bring strong foundational skills and a demonstrated ability to quickly learn and adapt.

CORE STRENGTHS:

"""
        for skill in list(cv['skills'].get('programming', []))[:3]:
            letter += f"✓ Strong proficiency in {skill}\n"
        
        letter += f"""
CURRENT DEVELOPMENT:

I am currently upskilling in the following areas to match your requirements:
"""
        for skill in gaps["missing_required_skills"][:3]:
            letter += f"• {skill} - Actively learning through hands-on projects\n"
        
        letter += f"""
I am a fast learner with a track record of mastering new technologies quickly. My {cv['experience_years']} years of experience have equipped me with strong problem-solving skills and the ability to deliver results under pressure.

I would appreciate the opportunity to discuss how my current skills and enthusiasm for continuous learning make me a strong candidate for this role.

Thank you for your consideration.

Best regards,
[Your Name]
"""
        return letter
    
    def _generate_future_interest_letter(self, analysis: Dict, learning_plan: Dict = None) -> str:
        """Generate letter expressing future interest"""
        job = analysis["job_info"]
        
        letter = f"""Dear Hiring Manager,

I am writing to express my interest in future opportunities at {job['company']}, specifically in roles similar to the {job['title']} position.

While I am currently building expertise in some key areas required for this role, I am highly motivated and following a structured learning plan to develop these skills. 

"""
        
        if learning_plan:
            letter += f"""MY DEVELOPMENT PLAN:

I am currently engaged in a {learning_plan['estimated_duration']} intensive learning program covering:
"""
            for item in learning_plan["levels"]["study"][:3]:
                letter += f"• {item['skill']}\n"
        
        letter += f"""
I am reaching out now to:
1. Express my strong interest in {job['company']}
2. Request to stay connected for future opportunities
3. Seek any advice on skill development priorities

I am targeting {datetime.now().month + 3} for actively applying to similar positions after completing my upskilling program.

Would it be possible to schedule a brief informational call to learn more about your team and future opportunities?

Thank you for your time and consideration.

Best regards,
[Your Name]
"""
        return letter
    
    def _generate_linkedin_message(self, analysis: Dict) -> str:
        """Generate LinkedIn connection request message"""
        job = analysis["job_info"]
        
        message = f"""Hi [Recruiter Name],

I recently came across the {job['title']} opportunity at {job['company']} and am very interested in this role. 

I have {analysis['cv_snapshot']['experience_years']}+ years of experience in related areas and would love to learn more about this position and {job['company']}.

Would you be open to connecting?

Best regards,
[Your Name]"""
        
        return message
    
    def _generate_followup_email(self, analysis: Dict) -> str:
        """Generate follow-up email template"""
        job = analysis["job_info"]
        
        email = f"""Subject: Following Up - {job['title']} Application

Dear Hiring Manager,

I wanted to follow up on my application for the {job['title']} position submitted on [DATE]. I remain very interested in this opportunity and wanted to reiterate my enthusiasm for joining {job['company']}.

Since submitting my application, I have:
- [Recent achievement/project]
- [Relevant skill development]
- [Industry contribution]

I would welcome the opportunity to discuss how my skills align with your team's needs. Are you available for a brief conversation this week?

Thank you for your time and consideration.

Best regards,
[Your Name]
[Your Contact Information]"""
        
        return email
    
    def _generate_networking_email(self, analysis: Dict) -> str:
        """Generate networking email to employees"""
        job = analysis["job_info"]
        
        email = f"""Subject: Seeking Insights About {job['company']}

Hi [Employee Name],

I hope this message finds you well. I am currently exploring opportunities in [field] and came across the {job['title']} position at {job['company']}.

I noticed your experience at {job['company']} on LinkedIn and would greatly appreciate any insights you could share about:
- The team culture and working environment
- Key skills that have proven valuable in your role
- Any advice for someone interested in joining the team

Would you have 15 minutes for a brief informational chat? I'd be grateful for any guidance.

Thank you for considering my request.

Best regards,
[Your Name]"""
        
        return email
    
    # ========== HELPER METHODS ==========
    
    def _extract_all_skills(self, text: str) -> Dict:
        """Extract all skills from text"""
        skills = defaultdict(list)
        text_lower = text.lower()
        
        skill_patterns = {
            "programming": ["python", "java", "javascript", "c++", "go", "rust", "scala", "r"],
            "frameworks": ["pytorch", "tensorflow", "react", "django", "flask", "spring", "angular"],
            "tools": ["docker", "kubernetes", "git", "jenkins", "terraform", "ansible"],
            "databases": ["postgresql", "mongodb", "mysql", "redis", "cassandra"],
            "cloud": ["aws", "azure", "gcp", "google cloud"]
        }
        
        for category, keywords in skill_patterns.items():
            for skill in keywords:
                if re.search(r'\b' + re.escape(skill) + r'\b', text_lower):
                    skills[category].append(skill)
        
        return dict(skills)
    
    def _extract_experience(self, text: str) -> int:
        """Extract years of experience"""
        patterns = [
            r'(\d+)\+?\s*years?\s+(?:of\s+)?experience',
            r'experience:\s*(\d+)',
            r'(\d+)\s*years?\s+in'
        ]
        years = []
        for pattern in patterns:
            matches = re.findall(pattern, text.lower())
            years.extend([int(m) for m in matches])
        return max(years) if years else 0
    
    def _extract_education(self, text: str) -> List[str]:
        """Extract education"""
        degrees = []
        if re.search(r'\b(phd|ph\.d|doctorate)\b', text.lower()):
            degrees.append("PhD")
        if re.search(r'\b(master|msc|m\.s|ma|mba)\b', text.lower()):
            degrees.append("Master's")
        if re.search(r'\b(bachelor|bsc|b\.s|ba)\b', text.lower()):
            degrees.append("Bachelor's")
        return degrees
    
    def _extract_certifications(self, text: str) -> List[str]:
        """Extract certifications"""
        certs = []
        cert_keywords = ["aws certified", "azure", "cfa", "pmp", "certified"]
        for cert in cert_keywords:
            if cert in text.lower():
                certs.append(cert)
        return certs
    
    def _extract_projects(self, text: str) -> List[str]:
        """Extract projects"""
        project_section = re.search(r'projects?:?\s*\n(.*?)(?:\n\n|\Z)', text, re.IGNORECASE | re.DOTALL)
        if project_section:
            return [p.strip() for p in project_section.group(1).split('\n') if p.strip()][:5]
        return []
    
    def _extract_soft_skills(self, text: str) -> List[str]:
        """Extract soft skills"""
        soft_skills = ["leadership", "communication", "teamwork", "problem solving", "analytical"]
        found = []
        text_lower = text.lower()
        for skill in soft_skills:
            if skill in text_lower:
                found.append(skill)
        return found
    
    def _extract_title(self, text: str) -> str:
        """Extract job title"""
        first_line = text.split('\n')[0].strip()
        return first_line[:100]
    
    def _extract_company(self, text: str) -> str:
        """Extract company name"""
        match = re.search(r'company:\s*([^\n]+)', text, re.IGNORECASE)
        return match.group(1).strip() if match else "Unknown"
    
    def _extract_required_skills(self, text: str) -> List[str]:
        """Extract required skills from job"""
        skills = []
        text_lower = text.lower()
        required_section = re.search(
            r'(?:required|must have).*?:(.*?)(?:preferred|nice to have|\Z)',
            text_lower, re.DOTALL
        )
        search_text = required_section.group(1) if required_section else text_lower
        
        all_skills = ["python", "java", "sql", "aws", "docker", "kubernetes", "pytorch", "tensorflow"]
        for skill in all_skills:
            if re.search(r'\b' + skill + r'\b', search_text):
                skills.append(skill)
        return skills
    
    def _extract_preferred_skills(self, text: str) -> List[str]:
        """Extract preferred skills"""
        skills = []
        preferred_section = re.search(
            r'(?:preferred|nice to have).*?:(.*?)(?:\Z)',
            text.lower(), re.DOTALL
        )
        if preferred_section:
            search_text = preferred_section.group(1)
            all_skills = ["react", "mongodb", "redis", "grafana", "airflow"]
            for skill in all_skills:
                if skill in search_text:
                    skills.append(skill)
        return skills
    
    def _extract_required_experience(self, text: str) -> int:
        """Extract required experience"""
        patterns = [r'(\d+)\+?\s*years?\s+(?:of\s+)?experience', r'minimum\s+(\d+)\s+years?']
        years = []
        for pattern in patterns:
            matches = re.findall(pattern, text.lower())
            years.extend([int(m) for m in matches])
        return min(years) if years else 0
    
    def _extract_education_req(self, text: str) -> List[str]:
        """Extract education requirements"""
        reqs = []
        text_lower = text.lower()
        if re.search(r'\b(phd|doctorate)\b', text_lower):
            reqs.append("PhD")
        if re.search(r'\b(master|graduate)\b', text_lower):
            reqs.append("Master's")
        if re.search(r'\b(bachelor|undergraduate)\b', text_lower):
            reqs.append("Bachelor's")
        return reqs
    
    def _extract_cert_req(self, text: str) -> List[str]:
        """Extract certification requirements"""
        certs = []
        cert_keywords = ["aws certified", "azure certified", "cfa", "pmp"]
        for cert in cert_keywords:
            if cert in text.lower():
                certs.append(cert)
        return certs
    
    def _extract_keywords(self, text: str) -> List[str]:
        """Extract important keywords"""
        keywords = []
        important = ["machine learning", "ai", "data science", "cloud", "devops"]
        for kw in important:
            if kw in text.lower():
                keywords.append(kw)
        return keywords
    
    def _extract_responsibilities(self, text: str) -> List[str]:
        """Extract job responsibilities"""
        resp_section = re.search(
            r'responsibilit(?:ies|y):?\s*(.*?)(?:requirements|qualifications|\Z)',
            text, re.IGNORECASE | re.DOTALL
        )
        if resp_section:
            return [r.strip() for r in resp_section.group(1).split('\n') if r.strip()][:5]
        return []
    
    def _calculate_score(self, cv_data: Dict, job_data: Dict) -> Dict:
        """Calculate match score"""
        scores = {
            "required_skills": self._score_skills(cv_data['skills'], job_data['required_skills']),
            "preferred_skills": self._score_skills(cv_data['skills'], job_data['preferred_skills']),
            "experience": self._score_experience(cv_data['experience_years'], job_data['required_experience']),
            "education": self._score_education(cv_data['education'], job_data['education_required']),
            "certifications": self._score_certs(cv_data['certifications'], job_data['certifications_required']),
            "keywords": self._score_keywords(cv_data.get('keywords', []), job_data['keywords'])
        }
        
        total = sum(scores[k] * self.WEIGHTS[k] for k in scores)
        
        return {
            "total_score": round(total, 2),
            "category_scores": scores,
            "weights": self.WEIGHTS
        }
    
    def _score_skills(self, cv_skills: Dict, required: List[str]) -> float:
        """Score skills match"""
        if not required:
            return 100.0
        cv_flat = [s.lower() for skills in cv_skills.values() for s in skills]
        matches = sum(1 for skill in required if skill.lower() in cv_flat)
        return (matches / len(required)) * 100
    
    def _score_experience(self, cv_years: int, required_years: int) -> float:
        """Score experience match"""
        if required_years == 0:
            return 100.0
        if cv_years >= required_years:
            return 100.0
        return (cv_years / required_years) * 80
    
    def _score_education(self, cv_edu: List[str], required: List[str]) -> float:
        """Score education match"""
        if not required:
            return 100.0
        levels = {"Bachelor's": 1, "Master's": 2, "PhD": 3}
        cv_level = max([levels.get(e, 0) for e in cv_edu], default=0)
        req_level = max([levels.get(e, 0) for e in required], default=0)
        return 100.0 if cv_level >= req_level else 70.0
    
    def _score_certs(self, cv_certs: List[str], required: List[str]) -> float:
        """Score certifications match"""
        if not required:
            return 100.0
        matches = sum(1 for cert in required if cert.lower() in [c.lower() for c in cv_certs])
        return (matches / len(required)) * 100
    
    def _score_keywords(self, cv_keywords: List[str], job_keywords: List[str]) -> float:
        """Score keywords match"""
        if not job_keywords:
            return 100.0
        matches = sum(1 for kw in job_keywords if kw.lower() in [k.lower() for k in cv_keywords])
        return (matches / len(job_keywords)) * 100
    
    def _identify_gaps(self, cv_data: Dict, job_data: Dict) -> Dict:
        """Identify skill gaps"""
        cv_flat = [s.lower() for skills in cv_data['skills'].values() for s in skills]
        
        return {
            "missing_required_skills": [s for s in job_data['required_skills'] if s.lower() not in cv_flat],
            "missing_preferred_skills": [s for s in job_data['preferred_skills'] if s.lower() not in cv_flat],
            "experience_gap": max(0, job_data['required_experience'] - cv_data['experience_years']),
            "education_gap": self._calc_edu_gap(cv_data['education'], job_data['education_required'])
        }
    
    def _calc_edu_gap(self, cv_edu: List[str], required: List[str]) -> str:
        """Calculate education gap"""
        if not required:
            return "None"
        levels = {"Bachelor's": 1, "Master's": 2, "PhD": 3}
        cv_level = max([levels.get(e, 0) for e in cv_edu], default=0)
        req_level = max([levels.get(e, 0) for e in required], default=0)
        return "None" if cv_level >= req_level else f"Need {required[0]}"
    
    def _generate_recommendations(self, gaps: Dict, score_result: Dict) -> List[Dict]:
        """Generate recommendations"""
        recs = []
        
        if gaps["missing_required_skills"]:
            recs.append({
                "priority": "CRITICAL",
                "category": "Required Skills",
                "action": f"Learn: {', '.join(gaps['missing_required_skills'][:3])}",
                "timeline": "2-4 weeks per skill"
            })
        
        if gaps["experience_gap"] > 0:
            recs.append({
                "priority": "HIGH",
                "category": "Experience",
                "action": f"Gain {gaps['experience_gap']} more years or build portfolio projects",
                "timeline": "Build 3-5 projects"
            })
        
        return recs
    
    def _get_resources(self, skill: str) -> Dict:
        """Get learning resources for skill"""
        skill_lower = skill.lower()
        return self.LEARNING_RESOURCES.get(skill_lower, self.LEARNING_RESOURCES["default"])
    
    def _get_prerequisites(self, skill: str) -> List[str]:
        """Get prerequisites for a skill"""
        prereqs = {
            "pytorch": ["python", "numpy", "basic ML"],
            "kubernetes": ["docker", "linux", "networking"],
            "aws": ["cloud concepts", "networking"],
            "react": ["javascript", "html", "css"]
        }
        return prereqs.get(skill.lower(), ["Basic programming knowledge"])
    
    def _generate_learning_path(self, skill: str) -> List[str]:
        """Generate step-by-step learning path"""
        return [
            f"1. Learn fundamentals of {skill}",
            f"2. Complete {skill} tutorials",
            f"3. Build 2-3 small projects",
            f"4. Study best practices",
            f"5. Contribute to open source or build portfolio project"
        ]
    
    def _generate_practice_goals(self, skill: str) -> List[str]:
        """Generate practice goals"""
        return [
            f"Complete 10+ {skill} coding challenges",
            f"Build 2 medium-complexity projects",
            f"Contribute to 1 open source project using {skill}",
            f"Write technical blog post about {skill}"
        ]
    
    def _check_certification(self, skill: str) -> bool:
        """Check if certification available"""
        cert_skills = ["aws", "azure", "gcp", "kubernetes", "terraform", "docker"]
        return skill.lower() in cert_skills
    
    def _generate_weekly_schedule(self, plan: Dict) -> List[Dict]:
        """Generate weekly schedule"""
        weeks = []
        for week in range(1, 13):
            weeks.append({
                "week": week,
                "focus": self._get_week_focus(week, plan),
                "hours": 15,
                "activities": self._get_week_activities(week, plan)
            })
        return weeks
    
    def _get_week_focus(self, week: int, plan: Dict) -> str:
        """Get focus for specific week"""
        if week <= 4:
            return f"Foundation: {plan['levels']['study'][0]['skill'] if plan['levels']['study'] else 'Core skills'}"
        elif week <= 8:
            return "Skill Development & Practice"
        else:
            return "Advanced Topics & Portfolio Building"
    
    def _get_week_activities(self, week: int, plan: Dict) -> List[str]:
        """Get activities for specific week"""
        if week <= 4:
            return ["Complete foundational course", "Daily practice (1hr)", "Build mini-project"]
        elif week <= 8:
            return ["Advanced course modules", "Coding challenges", "Medium project"]
        else:
            return ["Build portfolio project", "Mock interviews", "CV updates"]
    
    def _generate_milestones(self, plan: Dict) -> List[Dict]:
        """Generate milestones"""
        return [
            {"week": 4, "milestone": "Complete foundational learning", "deliverable": "3 mini-projects"},
            {"week": 8, "milestone": "Demonstrate intermediate proficiency", "deliverable": "2 portfolio projects"},
            {"week": 12, "milestone": "Job-ready skill level", "deliverable": "Complete portfolio + updated CV"}
        ]
    
    def _generate_action_items(self, analysis: Dict, learning_plan: Dict) -> List[Dict]:
        """Generate specific action items"""
        return [
            {"id": 1, "action": "Enroll in top-priority courses", "due": "Week 1", "status": "pending"},
            {"id": 2, "action": "Set up learning environment", "due": "Week 1", "status": "pending"},
            {"id": 3, "action": "Complete first project", "due": "Week 4", "status": "pending"},
            {"id": 4, "action": "Take beginner level tests", "due": "Week 4", "status": "pending"},
            {"id": 5, "action": "Build portfolio website", "due": "Week 8", "status": "pending"},
            {"id": 6, "action": "Take intermediate tests", "due": "Week 8", "status": "pending"},
            {"id": 7, "action": "Apply to target jobs", "due": "Week 12", "status": "pending"}
        ]
    
    def _count_total_skills(self) -> int:
        """Count total skills in skillset"""
        total = 0
        for category in self.master_skillset["technical_skills"].values():
            total += len(category)
        total += len(self.master_skillset["soft_skills"])
        total += len(self.master_skillset["certifications"])
        return total
    
    # ========== REPORTING METHODS ==========
    
    def generate_complete_report(self, analysis: Dict, learning_plan: Dict, 
                                 strategy: Dict, tests: Dict, letters: Dict) -> str:
        """Generate comprehensive report"""
        report = f"""
{'='*100}
                    COMPLETE JOB SEARCH & DEVELOPMENT REPORT
{'='*100}

Job: {analysis['job_info']['title']} at {analysis['job_info']['company']}
Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

{'='*100}
SECTION 1: JOB MATCH ANALYSIS
{'='*100}

Overall Match Score: {analysis['score']['total_score']}%

Category Breakdown:
"""
        for category, score in analysis['score']['category_scores'].items():
            weight = self.WEIGHTS[category]
            report += f"  • {category.replace('_', ' ').title():<25} {score:>6.1f}% (Weight: {weight:.0%})\n"
        
        report += f"""
Recommendation: {analysis['recommendations'][0]['action'] if analysis['recommendations'] else 'N/A'}

{'='*100}
SECTION 2: GAP ANALYSIS
{'='*100}

Missing Required Skills ({len(analysis['gaps']['missing_required_skills'])}):
"""
        for skill in analysis['gaps']['missing_required_skills'][:10]:
            report += f"  ❌ {skill}\n"
        
        report += f"""
Missing Preferred Skills ({len(analysis['gaps']['missing_preferred_skills'])}):
"""
        for skill in analysis['gaps']['missing_preferred_skills'][:5]:
            report += f"  ⚠️  {skill}\n"
        
        if analysis['gaps']['experience_gap'] > 0:
            report += f"\nExperience Gap: {analysis['gaps']['experience_gap']} years\n"
        
        report += f"""
{'='*100}
SECTION 3: LEARNING PLAN ({learning_plan['estimated_duration']})
{'='*100}

A) TO STUDY (New Content):
"""
        for i, item in enumerate(learning_plan['levels']['study'], 1):
            report += f"\n{i}. {item['skill'].upper()} [{item['priority']}]\n"
            report += f"   Time: {item['estimated_time']}\n"
            report += f"   Resources:\n"
            for res in item['resources'][:2]:
                report += f"     • {res}\n"
        
        report += f"\nB) TO PRACTICE (Strengthen Existing):\n"
        for i, item in enumerate(learning_plan['levels']['practice'], 1):
            report += f"\n{i}. {item['skill']}\n"
            report += f"   Practice: {', '.join(item['practice_activities'][:2])}\n"
        
        report += f"\nC) COURSES TO TAKE:\n"
        for i, item in enumerate(learning_plan['levels']['courses'][:5], 1):
            report += f"\n{i}. {item['skill']} - {item['recommended_courses'][0]}\n"
        
        report += f"""
{'='*100}
SECTION 4: IMPROVEMENT STRATEGY
{'='*100}

Timeline: {strategy['timeline']}
Current Score: {strategy['current_score']}%
Target Score: {strategy['target_score']}%

"""
        for phase in strategy['phases']:
            report += f"\n--- PHASE {phase['phase']}: {phase['name']} ({phase['duration']}) ---\n"
            report += f"Focus: {phase['focus']}\n"
            report += f"Expected Score Increase: +{phase['expected_score_increase']}%\n"
            report += f"Success Criteria:\n"
            for criteria in phase['success_criteria']:
                report += f"  ✓ {criteria}\n"
        
        report += f"""
{'='*100}
SECTION 5: SKILL ASSESSMENT TESTS
{'='*100}

Test Suite ID: {tests['test_suite_id']}
Skills Covered: {len(tests['skills_covered'])} skills

Test Structure:
  • Beginner Level: 10 questions, 60% pass score, 50 min
  • Intermediate Level: 15 questions, 70% pass score, 75 min
  • Advanced Level: 20 questions, 80% pass score, 100 min

Recommended Testing Schedule:
  Week 4: Beginner tests
  Week 8: Intermediate tests
  Week 12: Advanced tests

{'='*100}
SECTION 6: UPDATED SKILLSET
{'='*100}

New Skills to Add After Completion:
"""
        all_skills = set()
        for item in learning_plan['levels']['study']:
            all_skills.add(item['skill'])
        for item in learning_plan['levels']['practice']:
            all_skills.add(item['skill'])
            
        
        for skill in sorted(all_skills):
            report += f"  + {skill}\n"
        
        report += f"""
Current Total Skills: {self._count_total_skills()}
Projected Total Skills: {self._count_total_skills() + len(all_skills)}
"""
        # Add Reverse Workflow Progress section if applicable
        if self.state["mode"] == "reverse":
            report += f"""
{'='*100}
SECTION 6.5: REVERSE WORKFLOW PROGRESS
{'='*100}

Sprint-Based Development Status:

Current Sprint: {self.state['current_sprint']}
Total Learning Hours: {sum(s.get('total_hours', 0) for s in self.sprint_history if s.get('completed'))}h
Skills Mastered: {len(self.state['skills_mastered'])}
  → {', '.join(self.state['skills_mastered']) if self.state['skills_mastered'] else 'None yet'}

Projects Completed: {len(self.state['projects_completed'])}
"""
            for i, proj in enumerate(self.state['projects_completed'], 1):
                report += f"  {i}. Sprint {proj['sprint']}: {proj['goal']}\n"
                report += f"     URL: {proj['url']}\n"
            
            report += f"""
Quality Gates Passed: {', '.join(self.state['quality_gates_passed']) or 'None yet'}

Current Stage: {self.state['current_stage'].replace('_', ' ').title()}

Readiness Status:
  Brand Ready: {'âœ…' if self.state.get('brand_ready') else 'â³ In Progress'}
  Network Ready: {'âœ…' if self.state.get('network_ready') else 'â³ In Progress'}
  Application Ready: {'âœ…' if self.state.get('application_ready') else 'â³ In Progress'}
"""
        
        report += f"""
{'='*100}
SECTION 7: APPLICATION MATERIALS
{'='*100}

Generated Templates:
  ✓ Cover Letter (tailored)
  ✓ LinkedIn Connection Message
  ✓ Follow-up Email
  ✓ Networking Email

Application Strategy:
  1. Update CV with new skills and projects
  2. Create LinkedIn profile highlighting progress
  3. Reach out to company employees for networking
  4. Submit application with customized cover letter
  5. Follow up after 1 week

{'='*100}
NEXT STEPS
{'='*100}

IMMEDIATE (Week 1):
  1. Enroll in priority courses
  2. Set up development environment
  3. Start first learning module
  4. Connect with professionals on LinkedIn

SHORT-TERM (Weeks 2-8):
  5. Complete foundational learning
  6. Build 3-5 portfolio projects
  7. Take beginner and intermediate tests
  8. Contribute to open source

MEDIUM-TERM (Weeks 9-12):
  9. Complete advanced topics
  10. Build comprehensive portfolio
  11. Pass advanced level tests
  12. Update CV and apply to jobs

{'='*100}
REPORT END
{'='*100}
"""
        return report
    
    def export_all(self, job_id: str) -> str:
        """Export all data for a job analysis"""
        export_dir = self.data_dir / f"export_{job_id}"
        export_dir.mkdir(exist_ok=True)
        
        # Find analysis
        analysis = next((a for a in self.analyzed_jobs if a['job_id'] == job_id), None)
        if not analysis:
            return "Job analysis not found"
        
        # Generate all components
        learning_plan = self.create_learning_plan(analysis)
        strategy = self.create_improvement_strategy(analysis, learning_plan)
        tests = self.generate_skill_tests(analysis['gaps']['missing_required_skills'][:5])
        letters = self.generate_recruiter_letter(analysis, learning_plan)
        
        # Generate complete report
        report = self.generate_complete_report(analysis, learning_plan, strategy, tests, letters)
        
        # Save files
        files_created = []
        
        # 1. Complete report
        report_file = export_dir / "complete_report.txt"
        with open(report_file, 'w', encoding='utf-8') as f:
            f.write(report)
        files_created.append(str(report_file))
        
        # 2. Learning plan JSON
        plan_file = export_dir / "learning_plan.json"
        self._save_json(plan_file, learning_plan)
        files_created.append(str(plan_file))
        
        # 3. Strategy JSON
        strategy_file = export_dir / "improvement_strategy.json"
        self._save_json(strategy_file, strategy)
        files_created.append(str(strategy_file))
        
        # 4. Tests JSON
        tests_file = export_dir / "skill_tests.json"
        self._save_json(tests_file, tests)
        files_created.append(str(tests_file))
        
        # 5. Letters
        for letter_type, content in letters['templates'].items():
            letter_file = export_dir / f"{letter_type}.txt"
            with open(letter_file, 'w', encoding='utf-8') as f:
                f.write(content)
            files_created.append(str(letter_file))
        
        return f"✅ Exported {len(files_created)} files to {export_dir}"


 def main():
    """Interactive CLI for Advanced Job Engine"""
    print("\n" + "="*100)
    print("                        ADVANCED JOB ENGINE - COMPLETE SYSTEM")
    print("="*100)
    
    engine = AdvancedJobEngine()
    
    # Check for required libraries
    print("\n📚 Library Status:")
    print(f"   PDF Support (.pdf): {'✅' if PDF_AVAILABLE else '❌ (install PyPDF2)'}")
    print(f"   DOCX Support (.docx): {'✅' if DOCX_AVAILABLE else '❌ (install python-docx)'}")
    print(f"   Text Support (.txt): ✅")
    
    if not PDF_AVAILABLE or not DOCX_AVAILABLE:
        print("\n💡 To enable all formats, run:")
        if not PDF_AVAILABLE:
            print("   pip install PyPDF2")
        if not DOCX_AVAILABLE:
            print("   pip install python-docx")
    
    print("""
This system provides:
1. CV vs Job Analysis with Scoring (supports .txt, .pdf, .docx)
2. Gap Identification & Recommendations    
3. 3-Level Learning Plan (Study/Practice/Courses)
4. Strategic Improvement Plan
5. Multi-Level Skill Tests
6. Automatic Skillset Updates
7. Customized Recruiter Letters

Let's get started!
    """)
    
    # Demo mode
    print("\n[DEMO MODE]")
    
    demo_cv = """
John Smith
Data Scientist

Experience: 2 years in data analysis and machine learning

Skills:
- Programming: Python, SQL
- Tools: Pandas, Scikit-learn, Git
- Visualization: Matplotlib, Tableau

Education: Bachelor's in Computer Science

Projects:
- Built a customer churn prediction model
- Created data dashboard for sales analytics
    """
    
    demo_job = """
Senior Machine Learning Engineer

Company: TechCorp Inc.

Requirements:
- 5+ years experience in ML/AI
- Strong Python, PyTorch, TensorFlow
- AWS, Docker, Kubernetes
- MLOps experience
- Master's degree preferred

Skills:
- Deep Learning
- NLP
- Computer Vision
- Production ML systems
    """
    
    print("\n📊 Analyzing job opportunity...")
    
    # Step 1-2: Analyze job
    analysis = engine.analyze_job_complete(demo_cv, demo_job, 
                                          "Senior ML Engineer", "TechCorp")
    
    print(f"\n✅ Match Score: {analysis['score']['total_score']}%")
    print(f"📋 Missing {len(analysis['gaps']['missing_required_skills'])} required skills")
    
    # Step 3: Create learning plan
    print("\n📚 Creating learning plan...")
    learning_plan = engine.create_learning_plan(analysis)
    print(f"✅ {learning_plan['estimated_duration']} plan created")
    print(f"   - {len(learning_plan['levels']['study'])} skills to study")
    print(f"   - {len(learning_plan['levels']['practice'])} skills to practice")
    print(f"   - {len(learning_plan['levels']['courses'])} courses recommended")
    
    # Step 4: Create strategy
    print("\n🎯 Creating improvement strategy...")
    strategy = engine.create_improvement_strategy(analysis, learning_plan)
    print(f"✅ {len(strategy['phases'])} phase strategy")
    print(f"   Target: {strategy['target_score']}% match score")
    
    # Step 5: Generate tests
    print("\n📝 Generating skill assessment tests...")
    tests = engine.generate_skill_tests(analysis['gaps']['missing_required_skills'][:3])
    print(f"✅ 3-level tests for {len(tests['skills_covered'])} skills")
    
    # Step 6: Update skillset (simulated)
    print("\n🔄 Updating master skillset...")
    new_skills = ["PyTorch", "Docker", "Kubernetes"]
    update = engine.update_skillset(new_skills)
    print(f"✅ Added {len(update['updated_skills'])} new skills")
    print(f"   Total skills: {update['total_skills']}")
    
    # Step 7: Generate letters
    print("\n✉️  Generating recruiter letters...")
    letters = engine.generate_recruiter_letter(analysis, learning_plan)
    print(f"✅ Generated {len(letters['templates'])} templates:")
    for template_type in letters['templates'].keys():
        print(f"   - {template_type}")
    
    # Export everything
    print("\n📦 Exporting complete package...")
    export_result = engine.export_all(analysis['job_id'])
    print(export_result)
    
    # Generate and display complete report
    print("\n" + "="*100)
    print("COMPLETE REPORT")
    print("="*100)
    
    report = engine.generate_complete_report(analysis, learning_plan, strategy, tests, letters)
    print(report)
    
    print("\n✅ Demo completed successfully!")

      # Demonstrate reverse workflow
    print("\n" + "="*100)
    print("REVERSE WORKFLOW DEMONSTRATION")
    print("="*100)
    print("""
    The Reverse Workflow is an alternative approach focused on deep mastery before applying.

    PHILOSOPHY: Master skills FIRST, apply LAST

    Instead of:
    Analyze → Study 12 weeks → Apply

    You do:
    Analyze → Sprint 1 (2 weeks) → Sprint 2 (2 weeks) → ... → Quality Gates → Apply

    BENEFITS:
    ✓ Deeper skill mastery through iteration
    ✓ Real projects every 2 weeks
    ✓ Regular testing and validation
    ✓ Quality gates ensure readiness
    ✓ Stronger portfolio before applying
    """)
    
    print("\nStarting demo sprint...")
    sprint = engine.start_sprint(
        skills=["PyTorch", "Docker"],
        project_goal="Build and deploy ML model with Docker container"
    )
    
    print("\nSimulating daily learning logs...")
    engine.log_daily(3.5, ["PyTorch tensors", "Neural network basics"], "Completed PyTorch tutorial")
    engine.log_daily(4.0, ["Docker basics", "Containerization"], "Built first Docker image")
    engine.log_daily(3.0, ["Training models", "Model persistence"], "Trained first PyTorch model")
    
    print("\nEnding sprint with assessment...")
    result = engine.end_sprint(
        project_url="https://github.com/user/pytorch-docker-ml",
        test_scores={"PyTorch": 68, "Docker": 72}
    )
    
    print(f"\n✅ Sprint completed! Results:")
    print(f"   Skills mastered: {len(result.get('skills_mastered', []))}")
    print(f"   Quality gates: {', '.join(result.get('quality_gates_passed', [])) or 'None yet'}")
    
    # Show progress dashboard
    print("\nCurrent Progress:")
    engine.display_progress_dashboard()
    
    print("\n" + "="*100)
    print("                              SYSTEM FEATURES SUMMARY")
    print("="*100)
    print("""
✅ STEP 1-2: CV vs Job Analysis
   - Comprehensive skill matching
   - Experience and education scoring
   - Weighted scoring system (35% required skills, 20% experience, etc.)
   - Overall match percentage

✅ STEP 3: Three-Level Learning Plan
   A) TO STUDY: New content to learn from scratch
      - Detailed learning paths
      - Prerequisites identification
      - Resource recommendations
   
   B) TO PRACTICE: Existing skills to strengthen
      - Practice activities
      - Coding challenges
      - Real-world applications
   
   C) COURSES: Structured learning programs
      - Top-rated courses (Coursera, Udemy, etc.)
      - Certification options
      - Time estimates

✅ STEP 4: Strategic Improvement Plan
   - 12-week structured program
   - 3 phases: Foundation → Development → Mastery
   - Weekly schedules and milestones
   - Expected score improvements
   - Progress tracking system

✅ STEP 5: Multi-Level Skill Tests
   - Beginner Level: 10 questions, 60% pass
   - Intermediate Level: 15 questions, 70% pass
   - Advanced Level: 20 questions, 80% pass
   - Question templates for each skill
   - Assessment criteria and scoring

✅ STEP 6: Automatic Skillset Updates
   - Master skillset database
   - Automatic categorization
   - Skills tracking over time
   - Portfolio building

✅ STEP 7: Recruiter Letter Templates
   - Customized cover letters (3 types based on match score)
   - LinkedIn connection messages
   - Follow-up email templates
   - Networking email templates
   - All tailored to job description

📁 ALL DATA SAVED:
   - job_search_data/master_skillset.json
   - job_search_data/learning_progress.json
   - job_search_data/analyzed_jobs.json
   - job_search_data/skill_tests.json
   - job_search_data/export_[job_id]/ (complete package)
    """)
    
    print("\n" + "="*100)
    print("                                HOW TO USE")
    print("="*100)
    print("""
INTERACTIVE MODE:

from advanced_job_engine import AdvancedJobEngine

# Initialize
engine = AdvancedJobEngine()

# METHOD 1: Analyze from files (PDF, DOCX, TXT)
analysis = engine.analyze_from_files(
    cv_file="my_cv.pdf",
    job_file="job_description.docx",
    job_title="Senior ML Engineer",
    company="TechCorp"
)

# METHOD 2: Analyze from text
analysis = engine.analyze_job_complete(cv_text, job_text, "Job Title", "Company")

# Step 3: Create learning plan
learning_plan = engine.create_learning_plan(analysis)

# Step 4: Create improvement strategy
strategy = engine.create_improvement_strategy(analysis, learning_plan)

# Step 5: Generate tests
tests = engine.generate_skill_tests(missing_skills_list)

# Step 6: Update skillset after learning
engine.update_skillset(["PyTorch", "Docker", "AWS"], category="technical")

# Step 7: Generate recruiter letters
letters = engine.generate_recruiter_letter(analysis, learning_plan)

# Export everything
engine.export_all(analysis['job_id'])

# Generate complete report
report = engine.generate_complete_report(analysis, learning_plan, strategy, tests, letters)
print(report)

FILE FORMAT SUPPORT:
==================
✅ .txt  - Plain text files
✅ .pdf  - PDF documents (requires: pip install PyPDF2)
✅ .docx - Microsoft Word documents (requires: pip install python-docx)

INSTALLATION:
============
# Install all dependencies
pip install PyPDF2 python-docx

# Or install individually
pip install PyPDF2        # For PDF support
pip install python-docx   # For DOCX support

EXAMPLES:
========
# Analyze CV and job from PDF files
analysis = engine.analyze_from_files("cv.pdf", "job.pdf")

# Mix formats
analysis = engine.analyze_from_files("cv.docx", "job.txt")

# Read single document
cv_text = engine.read_document("my_cv.pdf")
job_text = engine.read_document("job_desc.docx")

# Generate complete report
report = engine.generate_complete_report(analysis, learning_plan, strategy, tests, letters)
print(report)
    """)
    
    print("\n" + "="*100)
    print("                         CUSTOMIZATION OPTIONS")
    print("="*100)
    print("""
1. ADJUST SCORING WEIGHTS:
   engine.WEIGHTS = {
       "required_skills": 0.40,  # Increase importance
       "experience": 0.25,
       # ... adjust as needed
   }

2. ADD CUSTOM LEARNING RESOURCES:
   engine.LEARNING_RESOURCES["your_skill"] = {
       "study": ["Resource 1", "Resource 2"],
       "practice": ["Practice 1", "Practice 2"],
       "courses": ["Course 1", "Course 2"]
   }

3. CUSTOMIZE TEST DIFFICULTY:
   engine.TEST_LEVELS["expert"] = {
       "questions": 25,
       "pass_score": 90
   }

4. MODIFY TIMELINE:
   - Edit strategy timeline (default: 12 weeks)
   - Adjust phase durations
   - Customize milestone schedule
    """)
    
    print("\n" + "="*100)
    print("                            WORKFLOW EXAMPLE")
    print("="*100)
    print("""
SCENARIO: Applying for Senior ML Engineer Role

Week 0 (Planning):
   1. Run job analysis → Get 55% match score
   2. Review gap analysis → Missing: PyTorch, Docker, Kubernetes, AWS
   3. Get learning plan → 12-week structured program
   4. Review improvement strategy → 3 phases to reach 80% match

Weeks 1-4 (Foundation):
   5. Study PyTorch fundamentals
   6. Complete Docker tutorials
   7. Build 3 mini-projects
   8. Take beginner-level tests → Pass with 70%
   9. Update skillset → Add "PyTorch (beginner)"

Weeks 5-8 (Development):
   10. Advanced PyTorch course
   11. Kubernetes hands-on practice
   12. Build 2 portfolio projects
   13. Take intermediate tests → Pass with 75%
   14. Update skillset → Upgrade to "intermediate"

Weeks 9-12 (Mastery):
   15. Build advanced ML project with full MLOps
   16. Deploy on AWS with Docker/K8s
   17. Take advanced tests → Pass with 85%
   18. Update skillset → Add all new skills
   19. Run job analysis again → Now 82% match!
   20. Generate recruiter letters
   21. Apply to job with confidence

Week 13+ (Application):
   22. Submit application with cover letter
   23. Connect on LinkedIn with hiring manager
   24. Send networking emails to team members
   25. Follow up after 1 week
   26. Interview preparation using test questions
   27. Land the job! 🎉
    """)

    print("\n" + "="*100)
    print("REVERSE WORKFLOW EXAMPLE (DETAILED)")
    print("="*100)
    print("""
    GOAL: Senior ML Engineer role requiring PyTorch, Docker, Kubernetes, AWS

    PHASE 1: FOUNDATION (Weeks 1-4, 2 sprints)
    ─────────────────────────────────────────
    Sprint 1 (PyTorch fundamentals):
    ├─ Day 1-7: Study PyTorch basics, tensors, autograd
    ├─ Day 8-14: Build image classifier project
    ├─ Test: PyTorch beginner (65% pass)
    └─ Result: 1 project, PyTorch skill mastered
  
    Sprint 2 (Docker fundamentals):
    ├─ Day 1-7: Docker basics, containers, images
    ├─ Day 8-14: Containerize PyTorch project
    ├─ Test: Docker beginner (70% pass)
    └─ Result: 2 projects, Docker skill mastered
  
    ✅ Quality Gate: FOUNDATION (2 projects, 65% score)

    PHASE 2: SKILL BUILDING (Weeks 5-12, 4 sprints)
    ─────────────────────────────────────────────
    Sprint 3 (Advanced PyTorch + Kubernetes intro):
    ├─ Build complex neural network architecture
    ├─ Deploy on Kubernetes cluster locally
    ├─ Tests: PyTorch intermediate (75%), K8s beginner (62%)
    └─ Result: 3 projects
  
    Sprint 4 (AWS basics + MLOps):
    ├─ Deploy model on AWS with Docker
    ├─ Set up CI/CD pipeline
    ├─ Tests: AWS beginner (68%), DevOps concepts (70%)
    └─ Result: 4 projects
  
    Sprint 5 (Integration project):
    ├─ End-to-end ML pipeline: PyTorch → Docker → K8s → AWS
    ├─ Monitoring and logging
    ├─ Tests: PyTorch advanced (82%), Docker intermediate (78%)
    └─ Result: 5 projects
    
    Sprint 6 (Production ML):
    ├─ Build production-grade ML system
    ├─ Scalability, reliability, monitoring
    ├─ Tests: All intermediate+ levels (75-85%)
    └─ Result: 6 projects
  
    ✅ Quality Gate: COMPETENCY (4 projects, 80% score)

    PHASE 3: MASTERY (Weeks 13-16, 2 sprints)
    ──────────────────────────────────────────
    Sprint 7 (Advanced architecture):
    ├─ Complex system design with all technologies
    ├─ Performance optimization
    ├─ Tests: Advanced levels (85%+)
    └─ Result: 7 projects
  
    Sprint 8 (Capstone project):
    ├─ Industry-grade project showcasing all skills
    ├─ Documentation, testing, deployment
    ├─ Tests: All advanced levels (85-90%)
    └─ Result: 8 projects, comprehensive portfolio
  
    ✅ Quality Gate: MASTERY (5+ projects, 90% score)

    PHASE 4: POSITIONING (Weeks 17-20)
    ───────────────────────────────────
    No sprints - focus on branding:
    ├─ Polish LinkedIn with all projects
    ├─ Organize GitHub portfolio professionally
    ├─ Write 3-5 technical blog posts
    ├─ Build personal website
    ├─ Network with 20+ professionals
    └─ Request informational interviews
  
    ✅ Quality Gate: APPLICATION READY (brand + network ready)

    PHASE 5: APPLICATION (Week 21+)
    ────────────────────────────────
    ├─ Re-analyze target jobs (now 90%+ match!)
    ├─ Generate customized application materials
    ├─ Apply to 10-15 target positions
    ├─ Follow up professionally
    ├─ Ace interviews with project portfolio
    └─ Land the job! 🎉

    TOTAL TIMELINE: 20-24 weeks
    RESULT: Deep mastery, strong portfolio, confident applications
    """)
    
    print("\n" + "="*100)
    print("                          KEY BENEFITS")
    print("="*100)
    print("""
✨ STRATEGIC:
   - Data-driven job targeting
   - Clear skill gap identification
   - Measurable progress tracking
   - ROI-focused learning

📚 COMPREHENSIVE:
   - End-to-end learning plan
   - Multiple learning modalities
   - Real-world practice emphasis
   - Portfolio building integrated

⚡ EFFICIENT:
   - Focused on job requirements
   - Prioritized skill development
   - Time-boxed learning phases
   - Quick wins + long-term growth

🎯 ACTIONABLE:
   - Specific action items
   - Weekly schedules
   - Clear success criteria
   - Ready-to-use templates

📊 MEASURABLE:
   - Regular skill assessments
   - Progress tracking
   - Score improvements
   - Portfolio growth

💼 PROFESSIONAL:
   - Customized application materials
   - Multiple letter templates
   - Networking strategies
   - Interview preparation
    """)
    
    print("\n" + "="*100)
    print("                        ADVANCED FEATURES")
    print("="*100)
    print("""
🔄 BATCH PROCESSING:
   - Analyze multiple jobs simultaneously
   - Compare opportunities side-by-side
   - Prioritize based on match scores
   - Identify common skill requirements

📈 PROGRESS TRACKING:
   - Historical skill development
   - Before/after comparisons
   - Learning velocity metrics
   - Portfolio evolution

🤝 COLLABORATION:
   - Export reports for mentors
   - Share learning plans
   - Team skill gap analysis
   - Hiring manager insights

🔍 ADVANCED ANALYTICS:
   - Skill demand trends
   - Learning effectiveness
   - Time-to-proficiency estimates
   - Career path recommendations

🎓 CERTIFICATION PLANNING:
   - Identify cert requirements
   - Study plan generation
   - Exam preparation resources
   - ROI analysis for certifications

💡 AI-POWERED (Future):
   - Personalized learning paths
   - Adaptive difficulty
   - Smart resource recommendations
   - Predictive success scoring
    """)
    
    print("\n" + "="*100)
    print("                          SUCCESS METRICS")
    print("="*100)
    print("""
Track your progress:

📊 Skill Score: 55% → 82% (+27%)
📁 Portfolio Projects: 1 → 6 (+5)
🎓 Certifications: 0 → 2 (+2)
💼 Job Applications: 0% success → Ready to apply
⏱️  Time to Job-Ready: 12 weeks
📝 Tests Passed: 0 → 9 (Beginner + Intermediate + Advanced)
🔧 Technical Skills: 8 → 18 (+10)
📧 Applications Sent: Customized & targeted
🤝 Network Connections: +20 relevant professionals
🎯 Interview Success: Prepared & confident
    """)
    
    print("\n" + "="*100)
    print("                      NEXT STEPS FOR YOU")
    print("="*100)
    print("""
1. 📄 Prepare your CV text
2. 🔍 Find job descriptions you're interested in
3. 🚀 Run the analysis for each job
4. 📊 Review match scores and gaps
5. 📚 Get your personalized learning plan
6. ⏰ Commit to 12-week development program
7. 📝 Take regular skill assessments
8. 🏗️  Build portfolio projects
9. 🔄 Update your skillset as you learn
10. ✉️  Use generated templates to apply
11. 🎯 Land your dream job!

Remember: Consistent daily effort beats sporadic intense study.
Aim for 15-20 hours per week of focused learning and practice.
    """)
    
    print("\n" + "="*100)
    print("Ready to transform your career? Let's go! 🚀")
    print("="*100 + "\n")

    # Interactive options
    print("\nWould you like to:")
    print("1. Analyze from text input")
    print("2. Analyze from files (.txt, .pdf, .docx)")
    print("3. View detailed demo report")
    print("4. Export demo results")
    print("5. Exit")
    
    choice = input("\nEnter choice (1-5): ").strip()
    
    if choice == "1":
        print("\n--- Analyze Your CV and Job (Text Input) ---")     
        print("\nPaste your CV (press Ctrl+D or Ctrl+Z on new line when done):")
        cv_lines = []
        try:
            while True:
                line = input()
                cv_lines.append(line)
        except EOFError:
            pass
        user_cv = "\n".join(cv_lines)
        
        print("\nPaste job description (press Ctrl+D or Ctrl+Z on new line when done):")
        job_lines = []
        try:
            while True:
                line = input()
                job_lines.append(line)
        except EOFError:
            pass
        user_job = "\n".join(job_lines)
        
        job_title = input("\nJob title: ")
        company = input("Company name: ")
        
        print("\n⏳ Analyzing... This may take a moment...")
        
        # Run complete analysis
        analysis = engine.analyze_job_complete(user_cv, user_job, job_title, company)
        learning_plan = engine.create_learning_plan(analysis)
        strategy = engine.create_improvement_strategy(analysis, learning_plan)
        tests = engine.generate_skill_tests(analysis['gaps']['missing_required_skills'][:5])
        letters = engine.generate_recruiter_letter(analysis, learning_plan)
        
        report = engine.generate_complete_report(analysis, learning_plan, strategy, tests, letters)
        print("\n" + report)
        
        save = input("\nSave complete package? (y/n): ")
        if save.lower() == 'y':
            result = engine.export_all(analysis['job_id'])
            print(result)
    
    elif choice == "2":
        
        print("\n--- Analyze from Files ---")
        print("\nSupported formats: .txt, .pdf, .docx")
        
        cv_file = input("\nPath to your CV file: ").strip().strip('"\'')
        job_file = input("Path to job description file: ").strip().strip('"\'')
        
        # Validate files exist
        if not Path(cv_file).exists():
            print(f"\n❌ Error: CV file not found: {cv_file}")
            return
        if not Path(job_file).exists():
            print(f"\n❌ Error: Job file not found: {job_file}")
            return
        
        job_title = input("Job title (optional): ").strip()
        company = input("Company name (optional): ").strip()
        
        print("\n⏳ Processing files...")
        
        try:
            # Run analysis from files
            analysis = engine.analyze_from_files(cv_file, job_file, job_title, company)
            
            print(f"\n✅ Analysis complete!")
            print(f"   Match Score: {analysis['score']['total_score']}%")
            print(f"   Missing Skills: {len(analysis['gaps']['missing_required_skills'])}")
            
            # Create learning plan and other materials
            print("\n📚 Creating comprehensive plan...")
            learning_plan = engine.create_learning_plan(analysis)
            strategy = engine.create_improvement_strategy(analysis, learning_plan)
            tests = engine.generate_skill_tests(analysis['gaps']['missing_required_skills'][:5])
            letters = engine.generate_recruiter_letter(analysis, learning_plan)
            
            # Generate report
            report = engine.generate_complete_report(analysis, learning_plan, strategy, tests, letters)
            print("\n" + report)
            
            save = input("\nSave complete package? (y/n): ")
            if save.lower() == 'y':
                result = engine.export_all(analysis['job_id'])
                print(result)
                print("\n📂 Files saved to: job_search_data/export_" + analysis['job_id'])
                
        except Exception as e:
            print(f"\n❌ Error: {str(e)}")
            import traceback
            traceback.print_exc()

    elif choice == "3":
        
        print("\n" + report)
    
    elif choice == "4":
        result = engine.export_all(analysis['job_id'])
        print("\n" + result)
        print("\n✅ All files exported successfully!")
        print("\nGenerated files:")
        print("  - complete_report.txt")
        print("  - learning_plan.json")
        print("  - improvement_strategy.json")
        print("  - skill_tests.json")
        print("  - cover_letter.txt")
        print("  - linkedin_message.txt")
        print("  - followup_email.txt")
        print("  - networking_email.txt")
    
    else:
        print("\n👋 Thank you for using Advanced Job Engine!")
        print("Good luck with your job search and skill development! 🚀\n")


if __name__ == "__main__":
    main()
